#!/usr/bin/env python3
"""Generate ENGINEERING_AUDIT_FULL.docx — comprehensive engineering audit of
the Bowtie Risk Analytics codebase."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import datetime

doc = Document()

# ── Style Setup ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    hs.font.name = "Calibri"

def add_table(headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    return t

def b(paragraph, text):
    """Add bold run."""
    r = paragraph.add_run(text)
    r.bold = True
    return r

def code(text):
    """Add code-formatted paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    return p

# ── Title Page ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n\nENGINEERING AUDIT\nFULL REPORT")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Bowtie Risk Analytics — Oil & Gas")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f"\nDate: {datetime.date.today().isoformat()}\n").font.size = Pt(12)
p.add_run("Audited by: Cross-Functional Senior Engineering Team\n").font.size = Pt(11)
p.add_run("(Senior SWE, Senior ML Engineer, Senior AI Engineer, Senior Bowtie/Cyber Risk Analyst)\n").font.size = Pt(10)
p.add_run("\nRepository: bowtie-risk-analytics-oilgas/\n").font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "Bowtie Risk Analytics is a Python pipeline, RAG retrieval system, and Streamlit dashboard "
    "for analyzing oil & gas incidents using the Bowtie risk methodology. It ingests public "
    "incident reports from CSB, BSEE, PHMSA, and TSB; extracts structured risk data via LLM; "
    "retrieves similar barrier failures via hybrid semantic search; and calculates barrier "
    "coverage metrics. The core question: \"Which barriers in this Bowtie are most likely to "
    "be weak or fail, and why?\""
)

doc.add_heading("Overall Codebase Health", level=2)
add_table(
    ["Layer", "Score /10", "Verdict"],
    [
        ["Data Pipeline (L0/L1/L2)", "8/10", "Mature. 739 incidents, 4,776 controls. Architecture frozen."],
        ["Schema V2.3 Model", "8/10", "97 fields, 5 validators, 17 sub-models. Comprehensive."],
        ["LLM Extraction", "7/10", "Policy-driven ladder, retry logic. Anthropic-only lock-in."],
        ["RAG Retrieval", "7/10", "4-stage hybrid + optional reranker. MRR 0.40, 40% recall gap."],
        ["Analytics Layer", "6/10", "Flatten + baseline functional. Engine.py is legacy V1."],
        ["Association Mining", "8/10", "45-family taxonomy. Production-quality normalization."],
        ["ML Modeling", "0/10", "Not built. src/modeling/ does not exist."],
        ["Streamlit Dashboard", "3/10", "Skeleton. 102 lines, V1 model, no RAG/V2.3 integration."],
        ["Test Suite", "8/10", "363 tests. Good coverage of pipeline, ingestion, RAG."],
        ["Documentation", "9/10", "Architecture freeze, pipeline contract, devlog all thorough."],
    ],
)

doc.add_heading("Critical Gaps", level=2)
doc.add_paragraph(
    "To answer 'Which barriers are most likely to fail, and why?', the system is missing: "
    "(1) Feature engineering module to derive ML features from controls + incidents CSVs; "
    "(2) ML training pipeline (LogReg + XGBoost) with stratified CV; "
    "(3) SHAP explainability layer for per-barrier reason codes; "
    "(4) RAG-to-LLM evidence wiring to generate narrative explanations; "
    "(5) A rebuilt Streamlit dashboard integrating predictions, SHAP, and RAG evidence."
)

doc.add_heading("Codebase Metrics", level=2)
add_table(
    ["Metric", "Value"],
    [
        ["Total Python source (src/)", "8,257 lines across 67 files"],
        ["Total Python tests (tests/)", "6,898 lines across 45 files"],
        ["Total Python scripts (scripts/)", "2,155 lines across 9 files"],
        ["Grand total Python", "17,310 lines"],
        ["Test count (pytest --co)", "363 tests collected"],
        ["Test coverage (estimated)", "~70-75% of src/ modules have corresponding tests"],
        ["Canonical data", "739 incidents, 4,776 controls (Schema V2.3)"],
        ["RAG corpus", "526 incidents, 3,253 barriers, 25 families"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: REPOSITORY INVENTORY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Repository Inventory", level=1)

doc.add_heading("2.1 Source File Tree with Line Counts", level=2)

src_files = [
    ("src/__init__.py", 1), ("src/analytics/__init__.py", 4), ("src/analytics/aggregation.py", 42),
    ("src/analytics/baseline.py", 112), ("src/analytics/build_combined_exports.py", 262),
    ("src/analytics/control_coverage_v0.py", 177), ("src/analytics/engine.py", 71),
    ("src/analytics/flatten.py", 107), ("src/app/__init__.py", 1), ("src/app/main.py", 102),
    ("src/app/utils.py", 38), ("src/corpus/__init__.py", 0), ("src/corpus/clean.py", 48),
    ("src/corpus/extract.py", 274), ("src/corpus/manifest.py", 76),
    ("src/extraction/__init__.py", 1), ("src/extraction/__main__.py", 49),
    ("src/extraction/extractor.py", 104), ("src/extraction/manifest.py", 65),
    ("src/extraction/normalize.py", 68), ("src/extraction/quality_gate.py", 80),
    ("src/extraction/runner.py", 151), ("src/ingestion/__init__.py", 0),
    ("src/ingestion/loader.py", 70), ("src/ingestion/manifests.py", 378),
    ("src/ingestion/normalize.py", 142), ("src/ingestion/pdf_text.py", 132),
    ("src/ingestion/source_ingest.py", 305), ("src/ingestion/structured.py", 500),
    ("src/ingestion/sources/__init__.py", 1), ("src/ingestion/sources/bsee.py", 187),
    ("src/ingestion/sources/bsee_discover.py", 168), ("src/ingestion/sources/csb.py", 303),
    ("src/ingestion/sources/csb_discover.py", 315), ("src/ingestion/sources/phmsa_discover.py", 244),
    ("src/ingestion/sources/phmsa_ingest.py", 144), ("src/ingestion/sources/tsb_discover.py", 199),
    ("src/ingestion/sources/tsb_ingest.py", 179), ("src/llm/__init__.py", 0),
    ("src/llm/anthropic_provider.py", 138), ("src/llm/base.py", 18),
    ("src/llm/model_policy.py", 51), ("src/llm/registry.py", 50), ("src/llm/stub.py", 97),
    ("src/models/__init__.py", 6), ("src/models/bowtie.py", 42), ("src/models/incident.py", 52),
    ("src/models/incident_v23.py", 429), ("src/nlp/loc_scoring.py", 305),
    ("src/pipeline.py", 886), ("src/prompts/__init__.py", 0), ("src/prompts/loader.py", 47),
    ("src/rag/__init__.py", 0), ("src/rag/config.py", 24),
    ("src/rag/context_builder.py", 70), ("src/rag/corpus_builder.py", 281),
    ("src/rag/embeddings/__init__.py", 0), ("src/rag/embeddings/base.py", 28),
    ("src/rag/embeddings/sentence_transformers_provider.py", 27),
    ("src/rag/rag_agent.py", 211), ("src/rag/reranker.py", 96),
    ("src/rag/retriever.py", 175), ("src/rag/vector_index.py", 96),
    ("src/validation/__init__.py", 0), ("src/validation/incident_validator.py", 28),
]

# Split into two tables for readability
mid = len(src_files) // 2
for chunk_start, chunk_end, label in [(0, mid, "Part A"), (mid, len(src_files), "Part B")]:
    add_table(
        ["File Path", "Lines"],
        [[f[0], str(f[1])] for f in src_files[chunk_start:chunk_end]],
    )
    doc.add_paragraph("")

doc.add_heading("2.2 Dependency Inventory", level=2)

add_table(
    ["Package", "Version Constraint", "Used By"],
    [
        ["streamlit", ">=1.30.0", "src/app/main.py — dashboard UI"],
        ["pydantic", ">=2.0.0", "src/models/, src/ingestion/manifests.py, structured.py"],
        ["pytest", ">=7.0.0", "tests/ — all test files"],
        ["python-dateutil", ">=2.8.0", "src/ingestion/manifests.py — date parsing"],
        ["requests", ">=2.28.0", "src/llm/anthropic_provider.py, all source scrapers"],
        ["pdfplumber", ">=0.10.0", "src/ingestion/pdf_text.py — primary PDF extraction"],
        ["PyMuPDF", ">=1.23.0", "src/extraction/extractor.py — fallback chain tier 1"],
        ["pdfminer.six", ">=20221105", "src/extraction/extractor.py — fallback chain tier 2"],
        ["pandas", ">=2.0.0", "src/analytics/, src/nlp/loc_scoring.py"],
        ["python-dotenv", ">=1.0.0", "src/pipeline.py — .env loading"],
        ["beautifulsoup4", ">=4.12.0", "src/ingestion/sources/tsb_discover.py, csb.py"],
        ["sentence-transformers", ">=2.2.0", "src/rag/embeddings/ — all-mpnet-base-v2"],
        ["faiss-cpu", ">=1.7.0", "src/rag/vector_index.py — FAISS IndexFlatIP"],
        ["numpy", ">=1.24.0", "src/rag/ — embedding arrays"],
        ["PyYAML", "(implicit)", "src/llm/model_policy.py — YAML config loading"],
    ],
)

doc.add_paragraph("")
p = doc.add_paragraph()
b(p, "Declared but not directly imported: ")
p.add_run("PyYAML is imported as 'yaml' but not listed explicitly in requirements.txt "
           "(it's a transitive dependency of other packages).")

doc.add_heading("2.3 Git Structure Notes", level=2)
doc.add_paragraph(
    "All data/ directories are gitignored except data/samples/ and data/evaluation/. "
    "The archive/ directory is gitignored. .env files are gitignored. "
    "The .codegraph/ directory exists for code graph indexing. "
    "To reproduce the full dataset, run the pipeline commands documented in CLAUDE.md."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: DATA MODEL AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Data Model Audit", level=1)

doc.add_heading("3.1 Legacy V1 Models (src/models/incident.py, bowtie.py)", level=2)
doc.add_paragraph("Status: LEGACY. Superseded by incident_v23.py. Still imported via src/models/__init__.py.")

add_table(
    ["Model", "File", "Lines", "Fields", "Validators", "Used By"],
    [
        ["Incident", "incident.py", "52", "17 fields (incident_id, date, location, facility_type, incident_type, severity, description, hazard, top_event, causes, consequences, prevention_barriers, mitigation_barriers, injuries, fatalities, environmental_impact, source)", "0", "src/app/main.py (dashboard), src/ingestion/loader.py, src/pipeline.py process_raw_files()"],
        ["Threat", "bowtie.py", "10", "3 fields (id, name, description)", "0", "src/analytics/engine.py"],
        ["Barrier", "bowtie.py", "22", "5 fields (id, name, description, type [Literal prevention/mitigation], effectiveness)", "0", "src/analytics/engine.py"],
        ["Consequence", "bowtie.py", "32", "4 fields (consequence_id, name, description, severity)", "0", "src/analytics/engine.py"],
        ["Bowtie", "bowtie.py", "43", "5 fields (hazard, top_event, threats, barriers, consequences)", "0", "src/analytics/engine.py, src/pipeline.py"],
    ],
)

doc.add_heading("3.2 Schema V2.3 Model (src/models/incident_v23.py — 429 lines, ACTIVE)", level=2)
doc.add_paragraph(
    "The canonical data model. 17 sub-models, 97 fields total, 5 validators (3 field_validator, 2 model_validator). "
    "All models use ConfigDict(strict=False) for flexible LLM output parsing. "
    "Backward-compat alias: IncidentV2_2 = IncidentV23 (line 429)."
)

v23_models = [
    ["SourceInfo", "12-29", "6", "doc_type (str='unknown'), url (Optional[str]), title (str='unknown'), date_published (Optional[str]), date_occurred (Optional[str]), timezone (Optional[str])", "0"],
    ["ContextInfo", "35-77", "4", "region (str='unknown'), operator (str='unknown'), operating_phase (str='unknown'), materials (list[str]=[])", "2: _stringify_operating_phase (list/dict/null -> str), _coerce_materials (dict/str/null -> list)"],
    ["EventInfo", "83-147", "7", "top_event (str='unknown'), incident_type (str='unknown'), costs (Optional[Union[str,int,float]]), actions_taken (list[str]=[]), summary (str=''), recommendations (list[str]=[]), key_phrases (list[str]=[])", "3: _remap_keys (model_validator: type->top_event, description->summary), _stringify_top_event, _normalize_costs"],
    ["HazardItem", "153-163", "3", "hazard_id (str), name (str), description (Optional[str])", "0"],
    ["ThreatItem", "165-175", "3", "threat_id (str), name (str), description (Optional[str])", "0"],
    ["ConsequenceItem", "177-188", "4", "consequence_id (str), name (str), description (Optional[str]), severity (Optional[str])", "0"],
    ["ControlPerformance", "194-225", "8", "barrier_status (Literal[active/degraded/failed/bypassed/not_installed/unknown]='unknown'), barrier_failed (bool=False), detection_applicable/mentioned (bool=False x2), alarm_applicable/mentioned (bool=False x2), manual_intervention_applicable/mentioned (bool=False x2)", "0"],
    ["ControlHuman", "227-246", "4", "human_contribution_value (Optional[str]), human_contribution_mentioned (bool=False), barrier_failed_human (bool=False), linked_pif_ids (list[str]=[])", "0"],
    ["ControlEvidence", "248-260", "2", "supporting_text (list[str]=[]), confidence (Literal[high/medium/low]='low')", "0"],
    ["ControlItem", "262-292", "12", "control_id (str), name (str='unknown'), side (Literal[prevention/mitigation]='prevention'), barrier_role (str='unknown'), barrier_type (Literal[engineering/administrative/ppe/unknown]='unknown'), line_of_defense (Literal[1st/2nd/3rd/recovery/unknown]='unknown'), lod_basis (Optional[str]), linked_threat_ids (list[str]=[]), linked_consequence_ids (list[str]=[]), performance (ControlPerformance), human (ControlHuman), evidence (ControlEvidence)", "0"],
    ["BowtieV2", "298-315", "4", "hazards (list[HazardItem]=[]), threats (list[ThreatItem]=[]), consequences (list[ConsequenceItem]=[]), controls (list[ControlItem]=[])", "0"],
    ["PeoplePifs", "321-334", "8", "competence_value/mentioned, fatigue_value/mentioned, communication_value/mentioned, situational_awareness_value/mentioned (4 x Optional[str]/bool pairs)", "0"],
    ["WorkPifs", "336-349", "8", "procedures, workload, time_pressure, tools_equipment (4 x value/mentioned pairs)", "0"],
    ["OrganisationPifs", "351-364", "8", "safety_culture, management_of_change, supervision, training (4 x value/mentioned pairs)", "0"],
    ["PifsInfo", "366-374", "3", "people (PeoplePifs), work (WorkPifs), organisation (OrganisationPifs)", "0"],
    ["NotesInfo", "380-392", "2", "rules (str='JSON output only...'), schema_version (str='2.3')", "0"],
    ["IncidentV23", "398-424", "7", "incident_id (str), source (SourceInfo), context (ContextInfo), event (EventInfo), bowtie (BowtieV2), pifs (PifsInfo), notes (NotesInfo)", "1: _remap_top_level (model_validator: moves misplaced top-level controls into bowtie)"],
]

add_table(
    ["Sub-Model", "Lines", "# Fields", "Fields (name, type, default)", "Validators"],
    v23_models,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: PIPELINE AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Pipeline Audit", level=1)

doc.add_heading("4.1 CLI Subcommands (src/pipeline.py — 886 lines)", level=2)

add_table(
    ["#", "Subcommand", "Lines", "Key Args", "Handler Function", "Reads From", "Writes To"],
    [
        ["1", "acquire", "552-577", "--csb-limit, --bsee-limit, --out, --download, --timeout, --append", "cmd_acquire()", "CSB/BSEE web", "data/raw/{csb,bsee}/manifest.csv + PDFs"],
        ["2", "extract-text", "580-591", "--manifest, --out", "cmd_extract_text()", "data/raw/*/pdf/", "data/raw/*/text/ + text_manifest.csv"],
        ["3", "process", "594-595", "(none)", "cmd_process()", "data/raw/, bowtie.json", "data/processed/"],
        ["4", "extract-structured", "598-635", "--text-dir, --out-dir, --manifest, --policy, --text-limit, --limit, --resume", "cmd_extract_structured()", "data/raw/*/text/", "data/structured/incidents/schema_v2_3/"],
        ["5", "convert-schema", "638-652", "--incident-dir (req), --out-dir (req)", "cmd_convert_schema()", "V2.2 JSON dir", "V2.3 JSON dir"],
        ["6", "schema-check", "655-663", "--incident-dir (default: schema_v2_3)", "cmd_schema_check()", "data/structured/incidents/schema_v2_3/", "(stdout/exit code)"],
        ["7", "quality-gate", "666-674", "--incident-dir (default: schema_v2_3)", "cmd_quality_gate()", "data/structured/incidents/schema_v2_3/", "(stdout)"],
        ["8", "extract-qc", "677-701", "--pdf-dir (req), --output-dir, --manifest, --force", "cmd_extract_qc()", "PDF directory", "archive/data/experimental_qc_extraction/"],
        ["9", "ingest-source", "704-737", "--source (req), --url-list, --input-pdf-dir, --output-root, --force, --timeout", "cmd_ingest_source()", "URL list or PDF dir", "data/raw/<source>/"],
        ["10", "ingest-phmsa", "740-760", "--csv-path (req), --output-dir, --manifest, --limit", "cmd_ingest_phmsa()", "PHMSA CSV", "data/raw/phmsa/"],
        ["11", "build-combined-exports", "763-787", "--incidents-dir, --output-dir", "cmd_build_combined_exports()", "data/structured/incidents/schema_v2_3/", "data/processed/flat_incidents_combined.csv + controls_combined.csv"],
        ["12", "discover-source", "790-828", "--source (req: csb/bsee/phmsa/tsb), --out, --limit, --base-url, --timeout, --sleep", "cmd_discover_source()", "Web scraping", "configs/sources/<source>/url_list.csv + metadata.csv"],
        ["13", "corpus-manifest", "830-834", "(none)", "cmd_corpus_manifest()", "data/corpus_v1/raw_pdfs/", "data/corpus_v1/manifests/corpus_v1_manifest.csv"],
        ["14", "corpus-clean", "836-844", "--dry-run", "cmd_corpus_clean()", "data/corpus_v1/", "Moves orphan JSONs to structured_json_noise/"],
        ["15", "corpus-extract", "846-868", "--policy, --delay, --text-limit", "cmd_corpus_extract()", "data/corpus_v1/ + raw text", "data/corpus_v1/structured_json/"],
    ],
)

doc.add_heading("4.2 Data Flow Diagram", level=2)
code("""L0: Raw Acquisition
  discover-source (CSB/BSEE/PHMSA/TSB) --> configs/sources/<src>/url_list.csv
  acquire / ingest-source               --> data/raw/<src>/pdf/ + manifest.csv
  extract-text                           --> data/raw/<src>/text/ + text_manifest.csv

L1: Structured Extraction
  extract-structured (LLM via model ladder) --> data/structured/incidents/schema_v2_3/*.json
  convert-schema (V2.2 -> V2.3 coercions)  --> data/structured/incidents/schema_v2_3/*.json
  schema-check (Pydantic validation)        --> exit code + report
  quality-gate (quality metrics)            --> stdout

L2: Analytics-Ready Exports
  build-combined-exports --> data/processed/flat_incidents_combined.csv (739 rows)
                         --> data/processed/controls_combined.csv (4,776 rows)

Scripts (out/):
  jsonaggregation.py     --> out/association_mining/incidents_aggregated.json
  jsonflattening.py      --> out/association_mining/incidents_flat.csv
  normalization.py       --> out/association_mining/normalized_df.csv""")

doc.add_heading("4.3 Ingestion Layer Detail", level=2)

add_table(
    ["Module", "Lines", "Purpose", "Status"],
    [
        ["ingestion/loader.py", "70", "Regex-based text-to-Incident V1 parser", "LEGACY"],
        ["ingestion/manifests.py", "378", "IncidentManifestRow, TextManifestRow, SourceManifestRow + merge logic", "ACTIVE"],
        ["ingestion/normalize.py", "142", "normalize_v23_payload(): V2.3 in-memory coercions (side, LoD, status)", "ACTIVE"],
        ["ingestion/pdf_text.py", "132", "pdfplumber text extraction + TextManifestRow generation", "ACTIVE"],
        ["ingestion/source_ingest.py", "305", "Generic URL-list/PDF-dir ingestion with SHA256 + resumability", "ACTIVE"],
        ["ingestion/structured.py", "500", "LLM extraction orchestrator, _parse_llm_json (3 strategies), StructuredManifestRow, quality gate", "ACTIVE"],
        ["sources/bsee.py", "187", "BSEE district page scraper + PDF download", "ACTIVE"],
        ["sources/bsee_discover.py", "168", "BSEE discovery (separated network/parse)", "ACTIVE"],
        ["sources/csb.py", "303", "CSB investigation scraper + PDF download (paginated, deny-list)", "ACTIVE"],
        ["sources/csb_discover.py", "315", "CSB discovery with 3-tier PDF URL scoring", "ACTIVE"],
        ["sources/phmsa_discover.py", "244", "PHMSA bulk data file discovery (CSV/XLSX, not PDFs)", "ACTIVE"],
        ["sources/phmsa_ingest.py", "144", "PHMSA CSV header inspection (skeleton, not fully implemented)", "SKELETON"],
        ["sources/tsb_discover.py", "199", "TSB pipeline report listing + HTML narrative extraction", "ACTIVE"],
        ["sources/tsb_ingest.py", "179", "TSB HTML report download + narrative extraction (not PDF-based)", "ACTIVE"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: ANALYTICS LAYER AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Analytics Layer Audit", level=1)

doc.add_heading("5.1 engine.py (71 lines) — LEGACY V1", level=2)
doc.add_paragraph(
    "Contains calculate_barrier_coverage() and identify_gaps(). Uses case-insensitive exact "
    "string matching of barrier names between Incident (V1) and Bowtie (V1) models. "
    "INCOMPATIBLE with Schema V2.3 (different field names, different model structure). "
    "Only used by the legacy 'process' subcommand and the dashboard."
)

doc.add_heading("5.2 flatten.py (107 lines) — CONTROLS_CSV_COLUMNS", level=2)
doc.add_paragraph("Defines the 16 canonical columns for the controls flat CSV:")
add_table(
    ["#", "Column Name", "Type", "Source"],
    [
        ["1", "incident_id", "str", "incident.incident_id"],
        ["2", "control_id", "str", "control.control_id"],
        ["3", "name", "str", "control.name"],
        ["4", "side", "Literal", "control.side (prevention/mitigation)"],
        ["5", "barrier_role", "str", "control.barrier_role"],
        ["6", "barrier_type", "Literal", "control.barrier_type"],
        ["7", "line_of_defense", "Literal", "control.line_of_defense"],
        ["8", "lod_basis", "Optional[str]", "control.lod_basis"],
        ["9", "linked_threat_ids", "str (CSV)", "control.linked_threat_ids joined"],
        ["10", "linked_consequence_ids", "str (CSV)", "control.linked_consequence_ids joined"],
        ["11", "barrier_status", "Literal", "control.performance.barrier_status"],
        ["12", "barrier_failed", "bool", "control.performance.barrier_failed"],
        ["13", "human_contribution_value", "Optional[str]", "control.human.human_contribution_value"],
        ["14", "barrier_failed_human", "bool", "control.human.barrier_failed_human"],
        ["15", "confidence", "Literal", "control.evidence.confidence"],
        ["16", "supporting_text_count", "int", "len(control.evidence.supporting_text)"],
    ],
)

doc.add_heading("5.3 build_combined_exports.py (262 lines)", level=2)
doc.add_paragraph(
    "Produces flat_incidents_combined.csv (739 rows) and controls_combined.csv (4,776 rows). "
    "4-tier source_agency resolution: (1) explicit agency/publisher field, (2) doc_type keyword match "
    "(e.g., 'bsee'->'BSEE'), (3) URL domain match (e.g., 'csb.gov'->'CSB'), (4) path segment match, "
    "(5) default 'UNKNOWN'. Controls CSV adds source_agency, provider_bucket, json_path to the 16 base columns."
)

doc.add_heading("5.4 control_coverage_v0.py (177 lines)", level=2)
doc.add_paragraph("GAP_STATUSES = {'failed', 'bypassed', 'not_installed', 'unknown', ''}.")
p = doc.add_paragraph()
b(p, "Coverage Score Formula: ")
p.add_run("coverage_score_v0 = (controls_active + 0.5 * controls_degraded) / controls_total")
doc.add_paragraph(
    "Produces 3 output CSVs: control_coverage_v0.csv (per-incident scores), "
    "control_gaps_v0.csv (individual gap rows), gap_rollups_v0.csv (aggregated by dimension)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: LLM LAYER AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("6. LLM Layer Audit", level=1)

doc.add_heading("6.1 Provider Architecture", level=2)
add_table(
    ["Module", "Lines", "Purpose", "Key Details"],
    [
        ["base.py", "18", "LLMProvider ABC", "Single abstract method: extract(prompt: str) -> str"],
        ["anthropic_provider.py", "138", "Claude API client", "API URL: https://api.anthropic.com/v1/messages; _RETRYABLE: {429,500,502,503,504}; Exponential backoff: 2^attempt; last_meta tracks provider/model/latency_ms/usage/stop_reason"],
        ["model_policy.py", "51", "YAML policy loader", "Frozen dataclass: provider, default_model, fallback_models, retries_per_model, promote_on (set of trigger strings)"],
        ["registry.py", "50", "Provider factory", "SUPPORTED_PROVIDERS = ('stub', 'anthropic'); _ENV_KEY_MAP = {'anthropic': 'ANTHROPIC_API_KEY'}; get_provider() dispatches by name"],
        ["stub.py", "97", "Test provider", "Returns hardcoded V2.3 JSON for STUB-001; incident_id, all 7 top-level sections present; schema_version='2.3'. Verified correct."],
    ],
)

doc.add_heading("6.2 Model Ladder (configs/model_policy.yaml)", level=2)
code("""provider: anthropic
default_model: claude-haiku-4-5-20251001
fallback_models:
  - claude-haiku-4-5-20251001      # Cheapest, try first
  - claude-sonnet-4-20250514       # Upgrade 1
  - claude-sonnet-4-5-20250929     # Upgrade 2
  - claude-sonnet-4-6              # Upgrade 3 (latest)
retries_per_model: 2               # 3 total attempts per model
promote_on: [timeout, rate_limit, empty_output, invalid_json, schema_validation_failed]""")

doc.add_heading("6.3 Prompt System", level=2)
doc.add_paragraph(
    "Template: assets/prompts/extract_incident.md (67 lines). Schema: assets/schema/incident_schema_v2_3_template.json "
    "(100 lines). Two substitution points: {{SCHEMA_TEMPLATE}} and {{INCIDENT_TEXT}}. "
    "Prompt includes 10 extraction rules, required field constraints, 5 enum definitions, and key name constraints. "
    "loader.py (47 lines) validates non-empty text and file existence before substitution."
)

doc.add_heading("6.4 Corpus Extraction Ladder (corpus/extract.py)", level=2)
doc.add_paragraph(
    "_run_model_ladder(): Iterates fallback_models list. For each model, creates AnthropicProvider(model=model_id, "
    "max_output_tokens=8192). Retries up to retries_per_model times. On failure, classifies error kind "
    "(rate_limit, timeout, invalid_json, schema_validation_failed, empty_output) and promotes to next model "
    "if kind is in promote_on set. run_corpus_extraction() orchestrates: load manifest, load text, build prompt, "
    "call ladder, normalize V2.3, validate, write JSON. Includes rate-limit delay estimation: "
    "tokens_est = len(text)/4; rate_wait = (tokens/30000) * 65."
)
doc.add_paragraph(
    "BUG: Validation failures are logged but invalid JSON is still written to disk (line 251). "
    "BUG: Truncation flag is returned but discarded (line 236)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: RAG SYSTEM AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("7. RAG System Audit", level=1)

doc.add_heading("7.1 Configuration (config.py — 24 lines)", level=2)
add_table(
    ["Constant", "Value", "Purpose"],
    [
        ["DEFAULT_EMBEDDING_MODEL", "all-mpnet-base-v2", "SentenceTransformer model"],
        ["EMBEDDING_DIMENSION", "768", "Vector dimensionality"],
        ["TOP_K_BARRIERS", "50", "Stage 1 candidate pool"],
        ["TOP_K_INCIDENTS", "20", "Stage 2 candidate pool"],
        ["TOP_K_FINAL", "10", "Final result count"],
        ["RRF_K", "60", "RRF denominator constant"],
        ["MAX_CONTEXT_CHARS", "8000", "Context builder char budget"],
        ["RERANKER_ENABLED", "True", "Toggle flag (UNUSED in code)"],
        ["RERANKER_MODEL", "cross-encoder/ms-marco-MiniLM-L-6-v2", "Cross-encoder model"],
        ["RERANKER_MAX_LENGTH", "512", "Cross-encoder max seq length"],
        ["RERANKER_BATCH_SIZE", "32", "Batch size"],
        ["TOP_K_RERANK", "30", "Reranker candidate pool"],
        ["FINAL_TOP_K", "10", "Post-rerank result count"],
    ],
)

doc.add_heading("7.2 4-Stage Hybrid Retrieval Pipeline (retriever.py — 175 lines)", level=2)
doc.add_paragraph(
    "Stage 1: Embed barrier_query, search barrier FAISS index with optional metadata mask "
    "(barrier_family, barrier_failed_human, PIF filters). Returns top-50.\n"
    "Stage 2: Embed incident_query, search incident FAISS index (unfiltered). Returns top-20.\n"
    "Stage 3: Intersection — retain only barriers whose parent incident appears in Stage 2 results.\n"
    "Stage 4: RRF ranking — score = 1/(k+barrier_rank) + 1/(k+incident_rank), k=60. Sort desc, return top-10."
)

p = doc.add_paragraph()
b(p, "RetrievalResult fields: ")
p.add_run("incident_id, control_id, barrier_family, barrier_failed_human, rrf_score, "
          "barrier_rank, incident_rank, barrier_sim_score, incident_sim_score, rerank_score (Optional)")

doc.add_heading("7.3 BARRIER_DOC_COLUMNS (corpus_builder.py — 26 columns)", level=2)
barrier_cols = [
    "incident_id", "control_id", "barrier_role_match_text", "barrier_family",
    "barrier_type", "side", "line_of_defense", "barrier_status", "barrier_failed",
    "barrier_failed_human", "human_contribution_value",
    "pif_competence", "pif_fatigue", "pif_communication", "pif_situational_awareness",
    "pif_procedures", "pif_workload", "pif_time_pressure", "pif_tools_equipment",
    "pif_safety_culture", "pif_management_of_change", "pif_supervision", "pif_training",
    "supporting_text", "confidence", "incident_summary",
]
add_table(["#", "Column"], [[str(i+1), c] for i, c in enumerate(barrier_cols)])

doc.add_heading("7.4 INCIDENT_DOC_COLUMNS (10 columns)", level=2)
incident_cols = [
    "incident_id", "incident_embed_text", "top_event", "incident_type",
    "operating_phase", "materials", "region", "operator", "summary", "recommendations",
]
add_table(["#", "Column"], [[str(i+1), c] for i, c in enumerate(incident_cols)])

doc.add_heading("7.5 RAGAgent.explain() — NO LLM Call", level=2)
doc.add_paragraph(
    "RAGAgent.explain() does NOT call an LLM. It returns ExplanationResult(context_text, results, metadata). "
    "Flow: (1) determine retrieval depth (30 if reranker, else top_k), (2) call HybridRetriever.retrieve(), "
    "(3) optionally call CrossEncoderReranker.rerank(), (4) build ContextEntry objects from barrier/incident "
    "metadata, (5) format as markdown via build_context() with MAX_CONTEXT_CHARS=8000 budget."
)

doc.add_heading("7.6 Baseline Performance (50-query benchmark)", level=2)
add_table(
    ["Metric", "Baseline (RRF)", "Reranked (+CE)", "Delta"],
    [
        ["Top-1 Accuracy", "0.30", "0.30", "+0.00"],
        ["Top-5 Hit Rate", "0.56", "0.56", "+0.00"],
        ["Top-10 Hit Rate", "0.62", "0.60", "-0.02"],
        ["MRR", "0.40", "0.42", "+0.01 (+3.1%)"],
        ["Avg Latency", "19 ms", "24 ms", "+5 ms"],
        ["Memory Overhead", "1328 MB", "1357 MB", "+29 MB"],
    ],
)
doc.add_paragraph(
    "Key finding: 20/50 queries are 'both miss' — expected barrier family not in top-10 for either system. "
    "Primary bottleneck is bi-encoder recall, not ranking quality. Reranker kept optional."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8: ASSOCIATION MINING AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Association Mining Audit", level=1)

doc.add_heading("8.1 Pipeline Overview", level=2)
add_table(
    ["Script", "Lines", "Input", "Output"],
    [
        ["jsonaggregation.py", "98", "data/structured/incidents/schema_v2_3/*.json", "out/association_mining/incidents_aggregated.json"],
        ["jsonflattening.py", "152", "incidents_aggregated.json", "out/association_mining/incidents_flat.csv (16 cols)"],
        ["event_barrier_normalization.py", "621", "incidents_flat.csv", "out/association_mining/normalized_df.csv (19 cols)"],
        ["smoke_test.py", "119", "Creates synthetic incident", "Validates full chain"],
    ],
)

doc.add_heading("8.2 4-Quadrant Barrier Family Taxonomy (45 families)", level=2)
doc.add_paragraph(
    "SIDE_MAP: left->prevention, right->mitigation. "
    "_QUADRANT_DISPATCH maps (barrier_level, barrier_type) to assignment function. "
    "ABBR_MAP expands 41 domain abbreviations (e.g., PSV->pressure safety valve, BOP->blowout preventer)."
)

add_table(
    ["Quadrant", "# Families", "Example Families"],
    [
        ["Prevention x Administrative", "10", "training, procedures, change_management, monitoring, regulatory_and_permits, hazard_analysis_prework_checks, operating_controls_and_limits, communication, planning, maintenance"],
        ["Prevention x Engineering", "5", "overpressurization_gas_discharge_gas_isolation, fluid_discharge_and_containment, prevention_of_ignition, detection_monitoring_alarms, mechanical_integrity"],
        ["Mitigation x Administrative", "10", "emergency_shutdown_isolation_depressurization, detection_monitoring_surveillance, active_intervention_to_stop_release, fire_response_firewatch_ignition_control, evacuation_muster_shelter_exclusion_access_control, medical_response_and_evacuation, environmental_response_cleanup_reporting, incident_command_coordination_and_comms, investigation_corrective_action_post_incident_verification, supervision_staffing_oversight"],
        ["Mitigation x Engineering", "18", "gas_detection_atmospheric_monitoring, alarms_general_alarm_pa, emergency_shutdown_isolation, emergency_disconnect_eds, well_control_barriers_kill, pressure_relief_blowdown_flare_disposal, ignition_source_control, active_fire_protection_firefighting, passive_fire_blast_protection, control_room_habitability_hvac_pressurization, emergency_power_backup_utilities, spill_containment_environmental_mitigation, chemical_release_scrubbing_neutralization, physical_protection_retention_restraints, emergency_escape_access_rescue_decon, structural_mechanical_integrity, remote_monitoring_intervention_subsea, marine_collision_avoidance"],
        ["Fallback", "2", "other_admin, other_engineering (for ppe/unknown types)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9: STREAMLIT DASHBOARD AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Streamlit Dashboard Audit", level=1)

doc.add_heading("9.1 main.py (102 lines) — SKELETON", level=2)
doc.add_paragraph(
    "Current state: Displays 4 KPI cards (Total Incidents, Avg Prevention Coverage, Avg Mitigation Coverage, "
    "Overall Coverage), an incident selector dropdown, and per-incident details (description, metadata table, "
    "barrier analysis with coverage % and gap list). "
    "Uses V1 Incident model (NOT V2.3). Reads pre-computed analytics from data/processed/. "
    "Has no RAG integration, no SHAP, no ML predictions, no search capability."
)

doc.add_heading("9.2 utils.py (38 lines)", level=2)
doc.add_paragraph(
    "load_data(processed_dir) returns (incidents_list, metrics_dict). "
    "Reads INC-*.json from processed/ and fleet_metrics.json. "
    "Returns empty lists/dicts if directory missing."
)

doc.add_heading("9.3 What's Missing", level=2)
add_table(
    ["Component", "Status"],
    [
        ["V2.3 schema integration", "Not connected — uses V1 Incident model"],
        ["RAG query interface", "Not built — no search/explain UI"],
        ["Barrier health predictions", "Not built — no ML model integration"],
        ["SHAP explanations", "Not built — no explainability layer"],
        ["Evidence narratives", "Not built — no RAG-to-LLM wiring"],
        ["Cross-agency comparison", "Not built — single-incident view only"],
        ["Barrier family taxonomy view", "Not built"],
        ["PIF heatmap", "Not built"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10: VALIDATION & NLP AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Validation & NLP Audit", level=1)

doc.add_heading("10.1 incident_validator.py (28 lines)", level=2)
doc.add_paragraph(
    "validate_incident_v23(payload: dict) -> (bool, list[str]): Wraps IncidentV23.model_validate(). "
    "Catches Pydantic ValidationError, formats errors as 'loc1 -> loc2: message'. "
    "Backward-compat alias: validate_incident_v2_2 = validate_incident_v23."
)

doc.add_heading("10.2 loc_scoring.py (305 lines) — LOC_v1 Frozen", level=2)
doc.add_paragraph(
    "Keyword-based Loss of Containment scoring. PRIMARY_LOC_TERMS: 8 terms (release, spill, leak, etc.). "
    "SECONDARY_LOC_TERMS: 2 terms (explosion, fire). HAZARDOUS_CONTEXT: 10 terms. "
    "Formula: loc_score = (primary*2) + (secondary*1) + hazardous. "
    "loc_flag = (primary>=1 AND hazardous>=1) OR (secondary>=1 AND hazardous>=2). "
    "Two entry points: run() for CSB manifest, run_with_extraction_manifest() for extraction-aware scoring."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11: TEST SUITE AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Test Suite Audit", level=1)

doc.add_paragraph(
    "Total: 45 test files, 363 test functions collected (verified via pytest --co). "
    "Architecture Freeze claims 300 tests (2026-03-04), README implies 362 (post-RAG merge). "
    "Current count is 363 — the 362 from README + 1 additional."
)

test_files = [
    ("test_aggregation.py", "analytics.aggregation", 1, 44),
    ("test_analytics.py", "analytics.engine", 2, 51),
    ("test_anthropic_provider.py", "llm.anthropic_provider", 7, 126),
    ("test_app_utils.py", "app.utils", 3, 67),
    ("test_baseline.py", "analytics.baseline", 8, 119),
    ("test_bowtie_models.py", "models.bowtie", 8, 98),
    ("test_build_combined_exports.py", "analytics.build_combined_exports", 13, 288),
    ("test_control_coverage_v0.py", "analytics.control_coverage_v0", 1, 46),
    ("test_corpus_clean.py", "corpus.clean", 4, 75),
    ("test_corpus_extract.py", "corpus.extract", 11, 380),
    ("test_corpus_manifest.py", "corpus.manifest", 8, 134),
    ("test_discover_sources.py", "ingestion.sources.* (CSB/BSEE/PHMSA)", 30, 530),
    ("test_extract_structured.py", "ingestion.structured", 20, 384),
    ("test_extraction_extractor.py", "extraction.extractor", 3, 42),
    ("test_extraction_manifest.py", "extraction.manifest", 4, 100),
    ("test_extraction_normalize.py", "extraction.normalize", 9, 56),
    ("test_extraction_pipeline_cli.py", "extraction CLI", 1, 18),
    ("test_extraction_quality_gate.py", "extraction.quality_gate", 8, 64),
    ("test_extraction_runner.py", "extraction.runner", 2, 92),
    ("test_flatten.py", "analytics.flatten", 7, 115),
    ("test_incident_validator.py", "validation.incident_validator", 23, 352),
    ("test_ingestion.py", "ingestion.loader", 2, 28),
    ("test_llm_provider.py", "llm.base + stub", 4, 31),
    ("test_loc_scoring_extraction_aware.py", "nlp.loc_scoring", 10, 183),
    ("test_manifest_merge.py", "ingestion.manifests (merge)", 17, 442),
    ("test_manifests.py", "ingestion.manifests (models)", 8, 174),
    ("test_models.py", "models.incident", 4, 48),
    ("test_pdf_text.py", "ingestion.pdf_text", 4, 89),
    ("test_phmsa_ingest.py", "sources.phmsa_ingest", 4, 57),
    ("test_pipeline_cli.py", "pipeline CLI", 10, 254),
    ("test_prompt_loader.py", "prompts.loader", 12, 91),
    ("test_rag_agent.py", "rag.rag_agent", 7, 303),
    ("test_rag_context_builder.py", "rag.context_builder", 4, 51),
    ("test_rag_corpus_builder.py", "rag.corpus_builder", 20, 350),
    ("test_rag_embeddings.py", "rag.embeddings", 6, 68),
    ("test_rag_integration.py", "RAG end-to-end", 2, 192),
    ("test_rag_reranker.py", "rag.reranker", 7, 185),
    ("test_rag_retriever.py", "rag.retriever", 9, 162),
    ("test_rag_vector_index.py", "rag.vector_index", 8, 67),
    ("test_registry.py", "llm.registry + JSON parsing", 10, 73),
    ("test_source_ingest.py", "ingestion.source_ingest", 10, 244),
    ("test_sources_bsee.py", "sources.bsee", 3, 90),
    ("test_sources_csb.py", "sources.csb", 8, 248),
    ("test_tsb_discover.py", "sources.tsb_discover", 12, 162),
    ("test_tsb_ingest.py", "sources.tsb_ingest", 2, 124),
]

add_table(
    ["Test File", "Module Tested", "# Tests", "Lines"],
    [[f[0], f[1], str(f[2]), str(f[3])] for f in test_files],
)

doc.add_heading("11.1 Coverage Gaps", level=2)
add_table(
    ["Source Module", "Has Tests?", "Notes"],
    [
        ["src/analytics/aggregation.py", "Yes (1 test)", "Minimal — no edge cases"],
        ["src/analytics/control_coverage_v0.py", "Yes (1 test)", "Single scenario only"],
        ["src/app/main.py", "No", "No Streamlit UI tests"],
        ["src/nlp/loc_scoring.py", "Partial", "run_with_extraction_manifest tested; run() not tested"],
        ["src/ingestion/sources/phmsa_ingest.py", "Minimal (4)", "Skeleton module, skeleton tests"],
        ["scripts/evaluate_retrieval.py", "No", "Evaluation script, not unit-tested"],
        ["scripts/generate_lessons_learned.py", "No", "Document generation, not tested"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12: BUG REGISTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Bug Register", level=1)

add_table(
    ["ID", "File:Line", "Severity", "Description", "Status"],
    [
        ["BUG-001", "pipeline.py:883-886", "CRITICAL", "Duplicate get_sources_root() with infinite recursion. Correct version at lines 6-9 returns Path('data/sources') as fallback. Duplicate at 883-886 calls itself recursively. Will cause RecursionError if configs/sources/ doesn't exist.", "OPEN — Must delete lines 883-886"],
        ["BUG-002", "pipeline.py:453-454", "LOW", "Duplicate mkdir call. Line 453 and 454 are identical: out_path.parent.mkdir(parents=True, exist_ok=True). Harmless but dead code.", "OPEN"],
        ["BUG-003", "build_combined_exports.py:158,221", "MEDIUM", "BOM encoding not enforced. Uses encoding='utf-8' instead of 'utf-8-sig' when reading V2.3 JSON files. Architecture Freeze mandates utf-8-sig for all V2.3 reads.", "OPEN"],
        ["BUG-004", "corpus/extract.py:251", "MEDIUM", "Validation failures still written to disk. If IncidentV23 validation fails, the JSON is still written. Should quarantine or fail.", "OPEN"],
        ["BUG-005", "corpus/extract.py:236", "LOW", "Truncation flag discarded. _run_model_ladder returns (data, truncated, model_used) but truncated is never used.", "OPEN"],
        ["BUG-006", "config.py:8", "LOW", "RERANKER_ENABLED flag defined as True but never read by any code. Reranker activation is controlled by presence of reranker parameter in RAGAgent constructor.", "OPEN (vestigial)"],
        ["BUG-007", "loc_scoring.py:25-41", "LOW", "Double-counting: 'explosion' and 'fire' appear in both SECONDARY_LOC_TERMS and HAZARDOUS_CONTEXT. A text with 'explosion fire' triggers loc_flag via secondary+hazardous path without any primary LOC term.", "BY DESIGN (but undocumented)"],
        ["BUG-008", "ingestion/structured.py:344-349", "MEDIUM", "Silent validation fallback. If model_validate fails but json.loads succeeds, unvalidated data is written without ERROR-level logging.", "OPEN"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13: ARCHITECTURE COMPLIANCE CHECK
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Architecture Compliance Check", level=1)

doc.add_heading("13.1 ARCHITECTURE_FREEZE_v1.md Compliance", level=2)
add_table(
    ["Rule", "Compliant?", "Evidence"],
    [
        ["Single canonical structured bucket (schema_v2_3/ only)", "YES", "No other incident dirs exist under data/structured/incidents/"],
        ["739 incidents, 739 unique IDs", "YES", "build-combined-exports produces 739 rows"],
        ["4,776 controls via get_controls()", "YES", "controls_combined.csv has 4,776 rows"],
        ["Layer isolation (L0->L1->L2)", "YES", "No reverse reads found in code"],
        ["BOM encoding (utf-8-sig read, utf-8 write)", "PARTIAL", "BUG-003: build_combined_exports.py uses utf-8 not utf-8-sig"],
        ["Provider bucketing in structured/", "YES", "Subdirs named by provider (anthropic/gemini/openai)"],
        ["Source bucketing in raw/", "YES", "Subdirs named by source (bsee/csb/phmsa/tsb)"],
        ["Single PDF directory per source", "YES", "pdf/ (not pdfs/)"],
        ["Manifests append-only", "YES", "merge functions upsert, never delete"],
        ["debug_llm_responses/ write-only", "YES", "No pipeline reads from it"],
        ["No archive/ reads", "YES", "No src/ code imports from archive/"],
        ["No rglob against structured/incidents/ root", "YES", "Code targets schema_v2_3/ directly"],
        ["Don't bypass get_controls()", "YES", "flatten.py reads bowtie.controls directly but does not claim to be get_controls()"],
        ["Don't store ML/RAG artifacts in structured/", "YES", "RAG data in data/evaluation/"],
        ["Don't write to out/ from pipeline.py", "YES", "Scripts write to out/, not pipeline"],
    ],
)

doc.add_heading("13.2 data_pipeline_contract_v1.md Compliance", level=2)
doc.add_paragraph(
    "All 10 invariants verified as compliant. Extension rules for RAG (data/rag/) are partially followed "
    "— RAG evaluation data is in data/evaluation/ rather than data/rag/, which is a minor deviation. "
    "All 10 'Do Not' rules are followed."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14: DEPENDENCY GRAPH
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Dependency Graph", level=1)

doc.add_heading("14.1 Internal Module Dependencies", level=2)
code("""pipeline.py
  -> ingestion/{loader, manifests, pdf_text, source_ingest, structured, normalize}
  -> ingestion/sources/{bsee, bsee_discover, csb, csb_discover, phmsa_*, tsb_*}
  -> analytics/{engine, aggregation, flatten, build_combined_exports}
  -> corpus/{manifest, clean, extract}
  -> extraction/runner
  -> models/{incident, bowtie, incident_v23}
  -> validation/incident_validator
  -> llm/{registry, model_policy}
  -> prompts/loader

rag/rag_agent.py
  -> rag/{retriever, reranker, context_builder, config}
  -> rag/embeddings/{base, sentence_transformers_provider}
  -> rag/vector_index
  -> rag/corpus_builder -> scripts/association_mining/event_barrier_normalization

corpus/extract.py
  -> llm/anthropic_provider (direct instantiation, bypasses registry)
  -> llm/model_policy
  -> ingestion/{structured._parse_llm_json, normalize}
  -> validation/incident_validator
  -> prompts/loader""")

doc.add_heading("14.2 External Dependencies", level=2)
add_table(
    ["Package", "Used By"],
    [
        ["requests", "llm/anthropic_provider, sources/bsee, csb, bsee_discover, csb_discover, tsb_*"],
        ["pdfplumber", "ingestion/pdf_text"],
        ["fitz (PyMuPDF)", "extraction/extractor (fallback tier 1)"],
        ["pdfminer.six", "extraction/extractor (fallback tier 2)"],
        ["pandas", "analytics/{baseline, build_combined_exports, control_coverage_v0, flatten}, nlp/loc_scoring"],
        ["numpy", "rag/{vector_index, retriever, reranker, corpus_builder, rag_agent, embeddings}"],
        ["faiss", "rag/vector_index"],
        ["sentence_transformers", "rag/embeddings/sentence_transformers_provider"],
        ["beautifulsoup4", "sources/{tsb_discover, tsb_ingest}"],
        ["yaml (PyYAML)", "llm/model_policy"],
        ["streamlit", "app/main"],
    ],
)

doc.add_heading("14.3 Circular Dependency Check", level=2)
doc.add_paragraph(
    "No circular dependencies found. The module graph is a DAG. "
    "One cross-boundary import: src/rag/corpus_builder.py imports from scripts/association_mining/event_barrier_normalization.py. "
    "This is a code smell — scripts/ should not be imported by src/."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 15: GAP ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("15. Gap Analysis: What's Built vs What's Needed", level=1)

add_table(
    ["Component", "Status", "Score", "What Exists", "What's Missing", "Effort"],
    [
        ["Data Pipeline (L0/L1/L2)", "Complete", "8/10", "15 CLI subcommands, 4 sources, manifest tracking, BOM encoding", "PHMSA/TSB ingestion incomplete; BOM bug in exports", "Low"],
        ["Schema V2.3 Model", "Complete", "8/10", "17 sub-models, 97 fields, 5 validators, backward-compat aliases", "No PIF consistency validation; no ID format validation", "Low"],
        ["LLM Extraction", "Complete", "7/10", "Policy-driven ladder, retry, 3 JSON parse strategies, quality gate", "Anthropic-only lock-in; invalid JSON still written; truncation flag unused", "Low"],
        ["RAG Retrieval", "Complete", "7/10", "4-stage hybrid pipeline, optional cross-encoder, 50-query benchmark", "40% recall gap; MRR 0.40; no BM25 hybrid; no query expansion", "Medium"],
        ["Association Mining", "Complete", "8/10", "45-family taxonomy, 41 abbreviations, 4-quadrant dispatch", "Scripts not importable as library; no unit tests for normalization accuracy", "Low"],
        ["Analytics Engine", "Partial", "5/10", "Flatten, baseline, combined exports, coverage_v0", "engine.py is V1-only legacy; no V2.3 analytics; no ML features", "Medium"],
        ["ML Modeling", "NOT BUILT", "0/10", "Nothing — src/modeling/ directory does not exist", "feature_engineering.py, train.py, explain.py, predict.py", "HIGH"],
        ["SHAP Explainability", "NOT BUILT", "0/10", "Nothing", "SHAP values, per-barrier reason codes, feature importance", "HIGH"],
        ["RAG->LLM Evidence", "NOT BUILT", "0/10", "RAGAgent returns context_text but no LLM call", "src/rag/explainer.py — wire RAG context + SHAP to LLM for narrative", "MEDIUM"],
        ["Streamlit Dashboard", "Skeleton", "3/10", "102-line V1 skeleton with 4 KPIs + incident explorer", "V2.3 integration, RAG search, predictions, SHAP, evidence", "HIGH"],
        ["FastAPI Backend", "NOT BUILT", "0/10", "Nothing", "REST API for predictions + RAG queries", "MEDIUM"],
        ["Next.js Frontend", "NOT BUILT", "0/10", "Nothing", "Modern React frontend", "HIGH"],
        ["Deployment", "NOT BUILT", "0/10", "Nothing", "Streamlit Community Cloud deployment", "LOW"],
        ["Documentation", "Thorough", "9/10", "Architecture freeze, pipeline contract, devlog, RAG reports", "No API docs, no user guide", "LOW"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 16: FEATURE ENGINEERING MODULE AUDIT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("16. Feature Engineering Module Audit", level=1)

doc.add_paragraph(
    "STATUS: NOT YET BUILT. The directory src/modeling/ does not exist. "
    "Neither src/modeling/feature_engineering.py nor tests/test_feature_engineering.py exist. "
    "This is the primary gap preventing the system from answering its core question."
)

doc.add_heading("16.1 Specification (from CLAUDE.md)", level=2)
doc.add_paragraph("Label derivation logic specified as:")
code("""barrier_did_not_perform = barrier_status in ('failed', 'degraded', 'not_installed', 'bypassed')
hf_contributed = barrier_failed_human == True
label = 1 if (barrier_did_not_perform and hf_contributed) else 0
# EXCLUDE rows where barrier_status == 'unknown'""")

doc.add_paragraph(
    "Features must be derived from controls_combined.csv joined with flat_incidents_combined.csv. "
    "barrier_status and barrier_failed_human must NOT be used as features (they define the label). "
    "12 PIF features (incident-level booleans) broadcast to all controls in that incident."
)

doc.add_heading("16.2 Available Feature Candidates", level=2)
add_table(
    ["Feature", "Type", "Source", "Notes"],
    [
        ["side", "Categorical (2)", "controls CSV", "prevention/mitigation"],
        ["barrier_type", "Categorical (4)", "controls CSV", "engineering/administrative/ppe/unknown"],
        ["line_of_defense", "Categorical (5)", "controls CSV", "1st/2nd/3rd/recovery/unknown"],
        ["barrier_role", "Text/Categorical", "controls CSV", "Free text — needs encoding or family assignment"],
        ["barrier_family", "Categorical (45)", "normalized_df.csv", "From association mining taxonomy"],
        ["source_agency", "Categorical (4)", "controls CSV", "CSB/BSEE/PHMSA/TSB"],
        ["supporting_text_count", "Numeric", "controls CSV", "Evidence quantity"],
        ["confidence", "Ordinal (3)", "controls CSV", "high/medium/low"],
        ["12 PIF _mentioned flags", "Binary (12)", "incidents CSV", "Broadcast from incident to all controls"],
        ["controls_per_incident", "Numeric", "Derived", "Count of controls per incident"],
        ["region", "Categorical", "incidents CSV", "Geographic region"],
        ["operating_phase", "Categorical", "incidents CSV", "startup/normal/shutdown/etc."],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 17: DATA QUALITY ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("17. Data Quality Assessment", level=1)

doc.add_heading("17.1 Expected Label Distribution", level=2)
doc.add_paragraph(
    "The label requires both conditions: barrier_did_not_perform (status in {failed, degraded, not_installed, bypassed}) "
    "AND barrier_failed_human == True. Based on typical industrial incident data, expect heavy class imbalance — "
    "most barriers are status='active' (worked as designed) and most have barrier_failed_human=False. "
    "Estimated positive rate: 5-15% of non-unknown controls. With 4,776 total controls and likely 10-20% unknown, "
    "the trainable dataset may be 3,800-4,300 rows with 200-600 positive labels."
)

doc.add_heading("17.2 PIF Sparsity Risk", level=2)
doc.add_paragraph(
    "12 PIF flags are at incident level, shared across all controls in that incident. "
    "With 739 incidents and 4,776 controls, average 6.5 controls per incident. "
    "PIF flags are binary (_mentioned) extracted by LLM. Risk: if most PIFs are False for most incidents, "
    "the PIF feature columns will be very sparse. However, the _mentioned flag indicates whether the PIF "
    "was discussed in the report, not whether it was a contributing factor — so sparsity may be lower "
    "than expected (reports often discuss multiple PIFs even when not causal)."
)

doc.add_heading("17.3 Confidence Field Interpretation", level=2)
doc.add_paragraph(
    "The 'confidence' field (high/medium/low) on each control is the LLM's extraction confidence — "
    "how certain the LLM was about the extracted fields, NOT domain expert confidence in barrier effectiveness. "
    "Prompt rules: high = direct quotes, medium = reasonable inference, low = unclear. "
    "This is a useful feature for ML: low-confidence extractions may correlate with sparse report coverage."
)

doc.add_heading("17.4 Unknown Exclusion Impact", level=2)
doc.add_paragraph(
    "Rows with barrier_status='unknown' are excluded from training per specification. "
    "Unknown status typically means the report didn't discuss this barrier's performance. "
    "Impact: reduces trainable dataset but removes noisy/uninformative rows. "
    "Must monitor the proportion of unknowns — if >30%, consider imputation strategies."
)

doc.add_heading("17.5 Controls per Incident Distribution", level=2)
doc.add_paragraph(
    "4,776 controls / 739 incidents = 6.5 average controls per incident. "
    "Distribution likely skewed: some incidents have 1-2 controls (minor events), "
    "some have 15-20+ (complex process safety events). The coverage_v0 module "
    "computes min/p50/p90/max distributions."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 18: RECOMMENDATIONS & ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("18. Recommendations & Roadmap", level=1)

doc.add_heading("18.1 Priority-Ordered Build List", level=2)

add_table(
    ["Priority", "Component", "Effort", "Description"],
    [
        ["P0 (Blocker)", "Fix BUG-001: Delete pipeline.py:883-886", "5 min", "Remove infinite recursion in duplicate get_sources_root()"],
        ["P0 (Blocker)", "Fix BUG-003: BOM encoding in build_combined_exports.py", "15 min", "Change encoding='utf-8' to 'utf-8-sig' at lines 158 and 221"],
        ["P1 (Critical)", "src/modeling/feature_engineering.py", "1-2 days", "Join controls + incidents, derive label, encode features, output feature matrix"],
        ["P1 (Critical)", "src/modeling/train.py", "1-2 days", "LogReg + XGBoost, stratified 5-fold CV, save to data/models/artifacts/"],
        ["P1 (Critical)", "src/modeling/explain.py", "1 day", "SHAP values + per-barrier reason codes"],
        ["P2 (High)", "src/modeling/predict.py", "0.5 day", "Inference on new/unseen controls"],
        ["P2 (High)", "src/rag/explainer.py", "1 day", "Wire RAG context + SHAP reasons to LLM API for evidence narrative"],
        ["P2 (High)", "Dashboard rebuild (src/app/)", "3-5 days", "Streamlit dashboard: V2.3, predictions, SHAP, RAG evidence"],
        ["P3 (Medium)", "RAG recall improvement", "2-3 days", "BM25 hybrid search, query expansion, domain-tuned embeddings"],
        ["P3 (Medium)", "Fix BUG-004: validation gate enforcement", "0.5 day", "Quarantine invalid JSONs instead of writing them"],
        ["P3 (Medium)", "Fix BUG-008: silent validation fallback", "0.5 day", "Add ERROR logging when model_validate fails"],
        ["P4 (Low)", "Deploy to Streamlit Community Cloud", "0.5 day", "Requires dashboard completion first"],
        ["P4 (Low)", "PHMSA/TSB ingestion completion", "2 days", "Complete skeleton implementations"],
        ["P4 (Low)", "Remove backward-compat aliases", "0.5 day", "Delete IncidentV2_2 and validate_incident_v2_2 aliases"],
    ],
)

doc.add_heading("18.2 Risk Register", level=2)
add_table(
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Class imbalance in ML target", "HIGH", "HIGH", "Use SMOTE/class weights; Precision@k evaluation; minimum 150-250 labeled examples"],
        ["PIF sparsity reduces feature utility", "MEDIUM", "MEDIUM", "Verify PIF distribution before training; consider PIF grouping (People/Work/Org aggregates)"],
        ["RAG 40% recall miss rate", "HIGH", "MEDIUM", "BM25 hybrid, query expansion, evaluate at 1000+ incidents; consider domain-tuned embeddings"],
        ["BOM encoding causes silent data corruption", "LOW", "HIGH", "Fix BUG-003; create read_v23_json() helper enforcing utf-8-sig everywhere"],
        ["Anthropic API cost at scale", "MEDIUM", "LOW", "Haiku-first ladder already in place; incremental extraction; prompt caching"],
        ["Pipeline infinite recursion", "LOW", "CRITICAL", "Fix BUG-001 immediately; only triggers if configs/sources/ missing"],
    ],
)

doc.add_heading("18.3 Architecture Decisions Needed", level=2)
add_table(
    ["Decision", "Options", "Recommendation"],
    [
        ["ML model storage location", "data/models/ vs. data/rag/models/", "data/models/artifacts/ per Architecture Freeze extension rules"],
        ["Feature matrix format", "CSV vs. Parquet vs. NumPy", "Parquet for efficient storage + pandas integration; CSV for debug"],
        ["SHAP integration pattern", "Per-prediction vs. precomputed", "Precompute for training set; on-demand for new predictions"],
        ["Evidence narrative generation", "Synchronous LLM call vs. pre-generate", "On-demand via explainer.py — keeps costs proportional to usage"],
        ["Dashboard architecture", "Streamlit-only vs. FastAPI+Streamlit", "Start Streamlit-only; add FastAPI if API consumers emerge"],
        ["Cross-boundary import fix", "Move normalization to src/ vs. keep in scripts/", "Refactor: extract shared normalization logic to src/analytics/normalization.py"],
    ],
)

# ── Footer ───────────────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading("Appendix: Audit Methodology", level=1)
doc.add_paragraph(
    "This audit was conducted by reading every line of every Python file in src/ (67 files, 8,257 lines), "
    "tests/ (45 files, 6,898 lines), and scripts/ (9 files, 2,155 lines), plus all documentation in docs/, "
    "all configuration files (requirements.txt, pyproject.toml, configs/model_policy.yaml, .gitignore), "
    "the prompt template (assets/prompts/extract_incident.md), and the schema template "
    "(assets/schema/incident_schema_v2_3_template.json). "
    "The audit was performed by parallel agent teams, each covering a specific module group, "
    "with findings cross-referenced against architecture documentation and pipeline contracts. "
    "Test count verified via 'pytest --co -q' (363 tests collected). "
    "Bug register compiled from code review findings across all modules."
)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = Path("docs/reports/ENGINEERING_AUDIT_FULL.docx")
out_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out_path))
print(f"Saved to {out_path}")
print(f"Sections: 18 + Appendix")
print("AUDIT COMPLETE")
