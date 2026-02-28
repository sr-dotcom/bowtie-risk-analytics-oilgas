"""Generate FINAL_AUDIT_v1.docx from audit findings."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import datetime


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def build_report() -> Document:
    doc = Document()

    # Title
    title = doc.add_heading("Bowtie Risk Analytics - Final Engineering Audit v1", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, f"Audit Date: {datetime.date.today().isoformat()}")
    add_para(doc, "Project: Bowtie Risk Analytics for Oil & Gas Barrier Risk Intelligence")
    add_para(doc, "Milestone: v1.0 Barrier Risk Intelligence (Phases 1-6 of 7 complete)")
    add_para(doc, "Stack: Python 3.12, FastAPI, Next.js 15, XGBoost, SHAP, FAISS, Pydantic v2")
    add_para(doc, "Tests: 455 Python (pytest) + 10 Frontend (vitest) = 465 total")
    doc.add_page_break()

    # ── SECTION 1: AI REFERENCE CLEANUP ──
    add_heading(doc, "Section 1: AI Reference Cleanup")
    add_para(doc, "Full scan of all .py, .ts, .tsx files for AI-generated markers, TODO/FIXME, placeholder text, hardcoded secrets, debug statements, and commented-out code.")

    add_heading(doc, "1.1 AI/LLM References", 2)
    add_para(doc, "No AI-generation watermarks found. All 'Claude' references are legitimate API product references (model names, provider config). Two comments use the phrase 'Claude\\'s discretion' which should be reworded for portfolio presentation:")
    add_table(doc, ["File", "Line", "Text", "Action"],
        [["src/api/schemas.py", "160", "# Error schema (Claude's discretion)", "Rewrite: '# Standard error body'"],
         ["src/modeling/explain.py", "48", "# Default bg size (Claude's discretion)", "Rewrite: '# Default background sample size: 200'"]])

    add_heading(doc, "1.2 TODO/FIXME/HACK", 2)
    add_para(doc, "Zero TODO, FIXME, or HACK comments found in any source file.")

    add_heading(doc, "1.3 Hardcoded Secrets", 2)
    add_para(doc, "No hardcoded API keys in source code. .env file exists on disk (not committed to git) with a live Anthropic API key. Action: Rotate the key as precaution.")

    add_heading(doc, "1.4 Debug Statements", 2)
    add_para(doc, "No stray print() in production library code. All print() calls are in CLI __main__ blocks (acceptable). One console.error in frontend BarrierForm.tsx error handler (acceptable in catch block).")

    add_heading(doc, "1.5 Commented-Out Code", 2)
    add_para(doc, "No commented-out code blocks found. Three section-divider comment blocks were flagged but are legitimate structural comments.")
    doc.add_page_break()

    # ── SECTION 2: CODE QUALITY ──
    add_heading(doc, "Section 2: Code Quality Audit (Senior SWE)")

    add_heading(doc, "2.1 Naming Conventions", 2)
    add_para(doc, "All public functions, classes, constants follow PEP 8 / camelCase (frontend) consistently. No violations.")

    add_heading(doc, "2.2 Type Hints", 2)
    add_para(doc, "98% coverage. 5 functions missing return type annotations (all in legacy/CLI code). 10 internal functions missing return types in ingestion/corpus modules. Frontend: zero 'any' types.")

    add_heading(doc, "2.3 Error Handling", 2)
    add_para(doc, "No bare except: clauses. 26 'except Exception' sites, all with logging or explicit fallback. One silent exception at src/ingestion/structured.py:307 should add logging.")

    add_heading(doc, "2.4 Import Organization", 2)
    add_para(doc, "One PEP 8 violation: src/pipeline.py has a function definition between import groups. 5 unused imports across legacy modules.")

    add_heading(doc, "2.5 Module Ratings", 2)
    add_table(doc, ["Module", "Rating", "Key Issues"],
        [["src/modeling/", "PRODUCTION-READY", "Strong typing, logging, docstrings"],
         ["src/rag/", "PRODUCTION-READY", "Clean architecture, tested"],
         ["src/api/", "PRODUCTION-READY", "lifespan missing return annotation"],
         ["src/llm/", "PRODUCTION-READY", "Clean ABC pattern"],
         ["src/ingestion/", "NEEDS POLISH", "Missing return types, one silent exception"],
         ["src/corpus/", "NEEDS POLISH", "Missing return types, no module docstring"],
         ["src/analytics/", "NEEDS POLISH", "Legacy modules lack docstrings"],
         ["src/app/", "NEEDS REFACTOR", "Streamlit skeleton pending rebuild"],
         ["src/pipeline.py", "NEEDS POLISH", "Import ordering, missing docstring"]])
    doc.add_page_break()

    # ── SECTION 3: ML ENGINEERING ──
    add_heading(doc, "Section 3: ML Engineering Audit (Senior ML Engineer)")
    add_para(doc, "Rating: PUBLISHABLE (with 2 medium fixes)", bold=True)

    add_heading(doc, "3.1 Feature Engineering", 2)
    add_para(doc, "No data leakage. Labels correctly derived from barrier_status formula (not barrier_failed column). GroupKFold on incident_id enforced. OrdinalEncoder fitted on full dataset (intentional - documented design decision). Unknown handling: -1 via unknown_value parameter.")

    add_heading(doc, "3.2 Model Training", 2)
    add_para(doc, "scale_pos_weight correctly computed from data (n_negative/n_positive). eval_metric in XGBoost constructor (3.0 API compliance). Metrics: F1-minority, MCC, Precision, Recall - no accuracy. random_state=42 throughout. Models retrained on full data after CV evaluation before saving artifacts.")

    add_heading(doc, "3.3 SHAP Explainability", 2)
    add_para(doc, "TreeExplainer on correct XGBoost models. Background: 200 samples. Values shape verified (n_samples x 18). TreeExplainer never serialized (recreated from model + .npy at load time).")
    add_para(doc, "ISSUE: float(expected_value) may fail if SHAP returns array for binary classification. Fix: use expected_value[1] for positive class.", bold=True)

    add_heading(doc, "3.4 PIF Ablation", 2)
    add_para(doc, "Methodology sound: identical GroupKFold splits for with/without comparison. Same scale_pos_weight per target. Result: Model 1 neutral, Model 2 improved +0.038. Advisory only (PIFs always stay).")

    add_heading(doc, "3.5 BarrierPredictor Pipeline", 2)
    add_para(doc, "Complete pipeline: raw dict -> encode via saved OrdinalEncoder -> predict_proba both models -> SHAP values both models -> PredictionResult. Risk thresholds from training percentiles (p60=0.858, p80=0.936).")

    add_heading(doc, "3.6 ML Issues", 2)
    add_table(doc, ["Severity", "File:Line", "Issue", "Fix"],
        [["MEDIUM", "src/modeling/predict.py:220,229", "float(expected_value) fails if SHAP returns array", "Use expected_value[1] if isinstance(expected_value, np.ndarray)"],
         ["MEDIUM", "src/modeling/explain.py:139", "Same expected_value issue in verification", "Same fix"],
         ["LOW", "src/modeling/train.py:219", "GroupKFold without stratification", "Consider StratifiedGroupKFold or fold-skip guard"],
         ["LOW", "scripts/generate_risk_thresholds.py:113", "No feature_names validation", "Add column order assertion"]])
    doc.add_page_break()

    # ── SECTION 4: RAG ENGINEERING ──
    add_heading(doc, "Section 4: AI/RAG Engineering Audit (Senior AI Engineer)")
    add_para(doc, "Rating: PRODUCTION-READY (with 2 medium fixes)", bold=True)

    add_heading(doc, "4.1 Disambiguation Fix", 2)
    add_para(doc, "PASS. _find_barrier_meta matches on control_id (not incident_id + barrier_family). control_id populated in retriever from barrier corpus CSV. Regression test confirms two same-family barriers return distinct metadata.")

    add_heading(doc, "4.2 Confidence Gate", 2)
    add_para(doc, "PASS. Hard gate at 0.65 on barrier_sim_score (not rrf_score). Below threshold returns 'No matching incidents found.' without LLM call. Empty results handled (default=0.0).")

    add_heading(doc, "4.3 BarrierExplainer", 2)
    add_para(doc, "SHAP reason codes in 'Model Analysis' prompt section. Prompt template well-structured with 3 substitution variables. ISSUE: No try/except around LLM extract() call - terminal API failures propagate uncaught.")

    add_heading(doc, "4.4 Citations", 2)
    add_para(doc, "Citation dataclass fully populated: incident_id, control_id, barrier_name, barrier_family, supporting_text (from corpus, first 500 chars), relevance_score.")

    add_heading(doc, "4.5 RAG Issues", 2)
    add_table(doc, ["Severity", "File:Line", "Issue", "Fix"],
        [["MEDIUM", "src/rag/explainer.py:91-92", "No try/except around LLM extract()", "Catch RuntimeError, return graceful fallback"],
         ["MEDIUM", "src/rag/reranker.py:44-50", "_find_meta() still matches on incident_id+barrier_family", "Add control_id to match criteria"]])
    doc.add_page_break()

    # ── SECTION 5: FRONTEND ──
    add_heading(doc, "Section 5: Frontend Audit (Senior SWE - Frontend)")

    add_heading(doc, "5.1 Component Ratings", 2)
    add_table(doc, ["Component", "Rating", "Notes"],
        [["lib/types.ts", "PRODUCTION-READY", "Complete, mirrors backend schemas"],
         ["lib/api.ts", "PRODUCTION-READY", "Clean, typed, throws on errors"],
         ["lib/riskScore.ts", "PRODUCTION-READY", "Pure function, tested"],
         ["context/BowtieContext.tsx", "PRODUCTION-READY", "Well-typed state"],
         ["diagram/BarrierNode.tsx", "PRODUCTION-READY", "Risk rings, hover states"],
         ["diagram/BowtieFlow.tsx", "PRODUCTION-READY", "nodeTypes at module scope"],
         ["diagram/TopEventNode.tsx", "PRODUCTION-READY", "Clean, simple"],
         ["diagram/layout.ts", "PRODUCTION-READY", "Dagre utility"],
         ["panel/DetailPanel.tsx", "PRODUCTION-READY", "Clean state routing"],
         ["panel/ShapWaterfall.tsx", "PRODUCTION-READY", "Tested, correct bars"],
         ["panel/EvidenceSection.tsx", "PRODUCTION-READY", "Loading/error/cache states"],
         ["panel/RiskScoreBadge.tsx", "PRODUCTION-READY", "Color-coded pill"],
         ["sidebar/BarrierForm.tsx", "NEEDS POLISH", "Missing aria-labels on form controls"],
         ["sidebar/constants.ts", "PRODUCTION-READY", "49 barrier families from encoder"]])

    add_heading(doc, "5.2 Accessibility Gap", 2)
    add_para(doc, "BarrierForm.tsx: 6 form controls (<input>, <select>, <textarea>) lack aria-label or htmlFor associations. This is the most significant frontend issue for production readiness.")
    doc.add_page_break()

    # ── SECTION 6: ARCHITECTURE ──
    add_heading(doc, "Section 6: Architecture Compliance")
    add_para(doc, "ALL 5 PHASE 1 BUGS VERIFIED FIXED:", bold=True)
    add_table(doc, ["Bug", "Location", "Status", "Verification"],
        [["BUG-001 infinite recursion", "pipeline.py:883-886", "FIXED", "Single get_sources_root() at line 6; file ends at 879"],
         ["BUG-003 BOM encoding", "flatten.py:90, build_combined_exports.py:158,221", "FIXED", "All three use utf-8-sig"],
         ["Duplicate mkdir", "pipeline.py:453-454", "FIXED", "Single mkdir call at 453"],
         ["control_coverage_v0 default", "control_coverage_v0.py:95", "FIXED", "Defaults to schema_v2_3"],
         ["control_id disambiguation", "rag_agent.py:203-211", "FIXED", "Matches on control_id"]])

    add_para(doc, "Architecture compliance: All model artifacts in data/models/artifacts/. No writes to data/structured/ or data/raw/ from modeling code. API code in src/api/. Frontend in frontend/.")
    add_para(doc, "One documentation gap: data/models/evaluation/ not listed in data_pipeline_contract_v1.md.")
    doc.add_page_break()

    # ── SECTION 7: REQUIREMENTS ──
    add_heading(doc, "Section 7: Requirements Satisfaction")
    add_para(doc, "33/36 requirements COMPLETE. 3 NOT STARTED (DEPLOY-01, DEPLOY-02, DEPLOY-03 - Phase 7).", bold=True)

    reqs = [
        ["DATA-01", "COMPLETE", "src/modeling/profile.py", "tests/test_profile.py", ""],
        ["DATA-02", "COMPLETE", "src/analytics/build_combined_exports.py", "tests/test_build_combined_exports.py", ""],
        ["DATA-03", "COMPLETE", "src/pipeline.py", "tests/test_pipeline_cli.py", ""],
        ["DATA-04", "COMPLETE", "src/modeling/profile.py", "tests/test_profile.py", ""],
        ["FEAT-01", "COMPLETE", "src/modeling/feature_engineering.py", "tests/test_feature_engineering.py", ""],
        ["FEAT-02", "COMPLETE", "src/modeling/feature_engineering.py", "tests/test_feature_engineering.py", ""],
        ["FEAT-03", "COMPLETE", "src/modeling/feature_engineering.py", "tests/test_feature_engineering.py", ""],
        ["FEAT-04", "COMPLETE", "src/modeling/feature_engineering.py", "tests/test_feature_engineering.py", ""],
        ["FEAT-05", "COMPLETE", "src/modeling/feature_engineering.py", "tests/test_feature_engineering.py", ""],
        ["MODEL-01", "COMPLETE", "src/modeling/train.py", "tests/test_train.py", ""],
        ["MODEL-02", "COMPLETE", "src/modeling/train.py", "tests/test_train.py", ""],
        ["MODEL-03", "COMPLETE", "src/modeling/train.py", "tests/test_train.py", ""],
        ["MODEL-04", "COMPLETE", "src/modeling/train.py", "tests/test_train.py", "No accuracy metric"],
        ["MODEL-05", "COMPLETE", "src/modeling/train.py", "tests/test_train.py", ""],
        ["SHAP-01", "COMPLETE", "src/modeling/explain.py", "tests/test_explain.py", ""],
        ["SHAP-02", "COMPLETE", "src/modeling/explain.py", "tests/test_explain.py", ""],
        ["SHAP-03", "COMPLETE", "src/modeling/predict.py", "tests/test_predict.py", "Separate dicts"],
        ["SHAP-04", "COMPLETE", "feature_engineering.py + predict.py", "tests/test_predict.py", "incident_context category"],
        ["SHAP-05", "COMPLETE", "src/modeling/explain.py", "tests/test_explain.py", "Advisory only"],
        ["RAG-01", "COMPLETE", "src/rag/explainer.py", "tests/test_rag_explainer.py", ""],
        ["RAG-02", "COMPLETE", "src/rag/explainer.py", "tests/test_rag_explainer.py", "Gate at 0.65"],
        ["RAG-03", "COMPLETE", "src/rag/rag_agent.py", "tests/test_rag_agent.py", "control_id match"],
        ["RAG-04", "COMPLETE", "src/rag/explainer.py", "tests/test_rag_explainer.py", ""],
        ["API-01", "COMPLETE", "src/api/main.py", "tests/test_api.py", ""],
        ["API-02", "COMPLETE", "src/api/main.py", "tests/test_api.py", "asyncio.to_thread"],
        ["API-03", "COMPLETE", "src/api/main.py", "tests/test_api.py", ""],
        ["API-04", "COMPLETE", "src/api/main.py", "tests/test_api.py", "Lifespan singleton"],
        ["API-05", "COMPLETE", "src/api/main.py", "tests/test_api.py", ""],
        ["UI-01", "COMPLETE", "frontend/components/diagram/", "vitest", "React Flow DAG"],
        ["UI-02", "COMPLETE", "frontend/components/panel/", "vitest", "SHAP waterfall"],
        ["UI-03", "COMPLETE", "frontend/components/panel/EvidenceSection.tsx", "vitest", "RAG evidence"],
        ["UI-04", "COMPLETE", "frontend/components/sidebar/BarrierForm.tsx", "vitest", "Input form"],
        ["UI-05", "COMPLETE", "frontend/components/diagram/BarrierNode.tsx", "vitest", "Color-coded nodes"],
        ["DEPLOY-01", "NOT STARTED", "-", "-", "Phase 7"],
        ["DEPLOY-02", "NOT STARTED", "-", "-", "Phase 7"],
        ["DEPLOY-03", "NOT STARTED", "-", "-", "Phase 7"],
    ]
    add_table(doc, ["Requirement", "Status", "Implementation", "Test", "Notes"], reqs)
    doc.add_page_break()

    # ── SECTION 8: PROFESSOR REQUIREMENTS ──
    add_heading(doc, "Section 8: Professor Requirements Check")
    add_table(doc, ["Requirement", "Status", "Evidence"],
        [["ONE prediction target (barrier health)", "SATISFIED", "Two models predict barrier failure likelihood (broad + HF-specific)"],
         ["2-3 models max", "SATISFIED", "2 architectures: LogReg baseline + XGBoost improved, x2 targets = 4 artifacts"],
         ["Human factors as INPUT FEATURES", "SATISFIED", "12 PIF _mentioned booleans are feature columns, not prediction targets"],
         ["RAG as EVIDENCE LAYER", "SATISFIED", "BarrierExplainer retrieves similar incidents, LLM generates evidence narrative"],
         ["Model explainability PRIMARY (SHAP)", "SATISFIED", "TreeExplainer on both XGBoost models, full SHAP vectors in API response"],
         ["Precision@k + qualitative rubric", "SATISFIED", "Precision, Recall, F1-minority, MCC per fold + mean/std"],
         ["150-250 min labeled examples", "EXCEEDED", "4,688 training-eligible controls (far exceeds minimum)"],
         ["Loss of Containment scenario", "SATISFIED", "Architecture, RAG corpus, demo scenario all LOC-scoped"],
         ["Demo accessible", "NEEDS WORK", "Local only (npm run dev + uvicorn). Phase 7 Docker needed."]])
    doc.add_page_break()

    # ── SECTION 9: TEAM FRAMEWORK ALIGNMENT ──
    add_heading(doc, "Section 9: Team Framework Alignment")

    add_heading(doc, "9.1 ML Framework", 2)
    add_para(doc, "Model 1 (broad barrier failure) and Model 2 (human factor sensitivity) both built as specified. PIF ablation study completed (advisory). Risk score 1-10 from training percentile distribution. SHAP explainability on both models separately.")

    add_heading(doc, "9.2 RAG Framework", 2)
    add_para(doc, "Four inputs wired: barrier query, incident query, SHAP reason codes, barrier metadata. Confidence gate at 0.65. Citations pinned to specific barriers. Disambiguation fix applied.")

    add_heading(doc, "9.3 UI Framework", 2)
    add_para(doc, "Single-page app with input form (sidebar), Bowtie diagram (center), detail panel (right). User inputs event + barriers, clicks Analyze, sees color-coded predictions + SHAP + evidence. Pre-filled demo scenario.")

    add_heading(doc, "9.4 Deviations from Teammate Specs", 2)
    add_para(doc, "Tailwind v3 instead of v4 (Node 18 constraint). Next.js 15 instead of 16 (same constraint). Vitest 2 instead of 4 (same). All documented as Rule 3 deviations.")
    doc.add_page_break()

    # ── SECTION 10: RECRUITER PERSPECTIVE ──
    add_heading(doc, "Section 10: Recruiter/Portfolio Perspective")

    add_heading(doc, "10.1 What Would Impress", 2)
    add_para(doc, "1. Full-stack ML pipeline: data profiling -> feature engineering -> training -> SHAP -> RAG -> API -> frontend. End-to-end in one repo.")
    add_para(doc, "2. 465 tests (455 Python + 10 vitest) with TDD discipline across all phases.")
    add_para(doc, "3. Architecture freeze document showing deliberate system design, not ad-hoc coding.")
    add_para(doc, "4. Real-world data: 739 incidents, 4,776 controls from CSB/BSEE government reports.")
    add_para(doc, "5. XGBoost + SHAP explainability: not just predictions but WHY a barrier is at risk.")
    add_para(doc, "6. RAG with hallucination guard: confidence gate prevents LLM fabrication on weak evidence.")
    add_para(doc, "7. Clean API design: FastAPI with Pydantic v2 schemas, asyncio.to_thread for blocking calls.")

    add_heading(doc, "10.2 What Would Concern", 2)
    add_para(doc, "1. README is only 36 lines - needs expansion with architecture diagram, screenshots, quick-start.")
    add_para(doc, "2. No Docker deployment yet (Phase 7 not started).")
    add_para(doc, "3. Legacy Streamlit skeleton (src/app/) still in codebase - should be removed or clearly marked.")
    add_para(doc, "4. 72.9% positive rate on Model 1 - heavily imbalanced majority-positive, unusual for failure prediction.")

    add_heading(doc, "10.3 Portfolio Recommendations", 2)
    add_para(doc, "HIGHLIGHT: ML pipeline design, SHAP waterfall visualization, RAG confidence gate, test coverage.")
    add_para(doc, "DOWNPLAY: Legacy pipeline code, Streamlit skeleton, deployment gap.")
    add_para(doc, "ADD: Architecture diagram (data flow), screenshot of Bowtie diagram, model performance table, GitHub badges (Python, tests passing, coverage).")
    doc.add_page_break()

    # ── SECTION 11: BUG REGISTER ──
    add_heading(doc, "Section 11: Bug Register")
    bugs = [
        ["BUG-R01", "src/modeling/predict.py:220", "MEDIUM", "float(expected_value) fails if SHAP returns numpy array for binary classification", "Use expected_value[1] if isinstance(ev, np.ndarray)"],
        ["BUG-R02", "src/modeling/explain.py:139", "MEDIUM", "Same expected_value issue in SHAP verification step", "Same fix as BUG-R01"],
        ["BUG-R03", "src/rag/explainer.py:91-92", "MEDIUM", "No try/except around LLM extract() call - uncaught RuntimeError on API failure", "Catch RuntimeError, return fallback ExplanationResult"],
        ["BUG-R04", "src/rag/reranker.py:44-50", "MEDIUM", "_find_meta() still matches on incident_id+barrier_family (not control_id)", "Add control_id to match criteria"],
        ["BUG-R05", "src/ingestion/structured.py:307", "LOW", "Silent break in retry loop with no logging", "Add logger.warning before break"],
        ["BUG-R06", "src/pipeline.py:6-15", "LOW", "get_sources_root() function defined between import groups (PEP 8 violation)", "Move function after all imports"],
        ["BUG-R07", "frontend/components/sidebar/BarrierForm.tsx", "LOW", "6 form controls lack aria-label/htmlFor accessibility attributes", "Add aria-label to all inputs/selects"],
        ["BUG-R08", "src/api/schemas.py:160", "INFO", "'Claude\\'s discretion' comment - remove for portfolio", "Rewrite comment"],
        ["BUG-R09", "src/modeling/explain.py:48", "INFO", "'Claude\\'s discretion' comment - remove for portfolio", "Rewrite comment"],
        ["BUG-R10", "README.md", "INFO", "Only 36 lines - insufficient for portfolio presentation", "Expand with architecture, screenshots, quick-start"],
        ["BUG-R11", "data_pipeline_contract_v1.md", "INFO", "data/models/evaluation/ directory not documented", "Update contract"],
    ]
    add_table(doc, ["ID", "File:Line", "Severity", "Description", "Fix"], bugs)
    doc.add_page_break()

    # ── SECTION 12: FINAL VERDICT ──
    add_heading(doc, "Section 12: Final Verdict")

    add_heading(doc, "12.1 Score Breakdown", 2)
    add_table(doc, ["Category", "Score", "Max", "Notes"],
        [["Code Quality", "17", "20", "Strong typing, clean architecture. Minor: import ordering, unused imports, missing docstrings on legacy modules"],
         ["ML Engineering", "18", "20", "Correct pipeline, GroupKFold, SHAP. Minor: expected_value latent bug, no stratified splits"],
         ["AI/RAG Integration", "13", "15", "Clean confidence gate, disambiguation fix. Minor: no LLM error handling, reranker still has old matching"],
         ["Frontend", "13", "15", "Clean components, React Flow well-used. Minor: accessibility gaps in form, one oversized component"],
         ["Testing", "9", "10", "465 tests, TDD discipline, test isolation excellent. Minor: no test for risk_thresholds script"],
         ["Documentation", "7", "10", "Architecture freeze, data contract, CLAUDE.md excellent. README thin, DEVLOG sparse, data_pipeline_contract outdated"],
         ["Architecture", "9", "10", "Freeze respected, layer isolation clean. Minor: contract not updated for data/models/evaluation/"]])

    add_heading(doc, "TOTAL SCORE: 86/100", 2)
    add_para(doc, "GO FOR DEPLOYMENT (after Phase 7 + priority fixes)", bold=True)

    add_heading(doc, "12.2 Priority Fixes Before Deployment", 2)
    add_para(doc, "P0 (Before demo):")
    add_para(doc, "  1. Fix SHAP expected_value handling in predict.py and explain.py (BUG-R01, R02)")
    add_para(doc, "  2. Add try/except around LLM extract() in explainer.py (BUG-R03)")
    add_para(doc, "  3. Fix reranker _find_meta() to match on control_id (BUG-R04)")
    add_para(doc, "")
    add_para(doc, "P1 (Before portfolio submission):")
    add_para(doc, "  4. Expand README.md with architecture diagram, screenshots, quick-start guide")
    add_para(doc, "  5. Remove 'Claude\\'s discretion' comments (BUG-R08, R09)")
    add_para(doc, "  6. Add aria-labels to BarrierForm.tsx (BUG-R07)")
    add_para(doc, "  7. Complete Phase 7 (Docker Compose deployment)")
    add_para(doc, "")
    add_para(doc, "P2 (Nice to have):")
    add_para(doc, "  8. Fix import ordering in pipeline.py (BUG-R06)")
    add_para(doc, "  9. Remove unused imports (5 files)")
    add_para(doc, "  10. Add missing return type annotations (15 functions)")
    add_para(doc, "  11. Update data_pipeline_contract_v1.md for data/models/evaluation/")

    add_heading(doc, "12.3 Deployment Readiness", 2)
    add_para(doc, "Phases 1-6 are COMPLETE and VERIFIED. The application works end-to-end: data profiling -> feature engineering -> model training -> SHAP explainability -> RAG evidence -> FastAPI backend -> Next.js frontend.")
    add_para(doc, "Phase 7 (Docker Compose deployment) is the final step. The 4 P0 bugs are runtime risks that should be fixed before any live demo. The P1 items are portfolio presentation concerns.")
    add_para(doc, "Overall: This is a well-engineered, well-tested ML application with clear architecture, strong separation of concerns, and comprehensive explainability. Ready for deployment after the priority fixes.")

    return doc


if __name__ == "__main__":
    out_path = Path("docs/reports/FINAL_AUDIT_v1.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(str(out_path))
    print(f"Audit report written to: {out_path}")
    print(f"Size: {out_path.stat().st_size:,} bytes")
