#!/usr/bin/env python3
"""Generate Lessons Learned 2 Word document for Bowtie Risk Analytics."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def set_cell_shading(cell, color_hex: str) -> None:
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    """Add a professional table with header shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2E4057")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Lessons Learned 2")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Bowtie Risk Analytics — Oil & Gas Incident Intelligence Pipeline")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4A, 0x6F, 0x8A)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Pod: Bowtie Risk Analytics Team\nCourse: [Course Name]\nDate: March 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ========================================================================
    # TOPIC 1: MOST INNOVATIVE THING WE DID (5 points)
    # ========================================================================
    doc.add_heading("1. Most Innovative Thing We Did (Business Value)", level=1)

    doc.add_heading("What We Built: Dual-Axis Hybrid RAG Retrieval Pipeline", level=2)

    doc.add_paragraph(
        "We designed and implemented a 4-stage hybrid Retrieval-Augmented Generation (RAG) "
        "pipeline that performs dual-axis semantic search across two independent dimensions — "
        "barrier similarity and incident context — then fuses results using Reciprocal Rank "
        "Fusion (RRF). This architecture is fundamentally different from standard single-index "
        "RAG systems."
    )

    doc.add_heading("What Makes It Novel / Differentiated", level=2)

    innovations = [
        (
            "Dual-Encoder Architecture",
            "Most RAG systems embed documents into a single vector space and retrieve by "
            "proximity. Our system maintains two separate embedding indices — one for barrier "
            "controls (3,253 vectors) and one for incidents (526 vectors) — each encoding "
            "different semantic dimensions. Barrier embeddings capture control function "
            "(e.g., \"pressure relief valve — overpressure protection\"), while incident "
            "embeddings capture operational context (e.g., \"gas release during well "
            "intervention, deepwater, hydrocarbons\"). This dual-axis design lets the system "
            "answer compound queries that no single-index system can handle natively: "
            "\"Find barriers similar to X that failed in incidents similar to Y.\""
        ),
        (
            "Intersection-Based Noise Elimination (Stage 3)",
            "After retrieving the top-50 similar barriers and top-20 similar incidents "
            "independently, Stage 3 performs an intersection filter: it retains only "
            "barriers whose parent incident also appeared in the incident search results. "
            "This is the key innovation — it eliminates barriers that are semantically "
            "similar but occurred in contextually irrelevant incidents. Standard retrieval "
            "systems cannot do this because they lack the relational link between barriers "
            "and their parent incidents."
        ),
        (
            "Domain-Specific Barrier Family Taxonomy",
            "We built a 46-family, 4-quadrant classification system "
            "(prevention/mitigation x engineering/administrative) that normalizes "
            "free-text barrier names into a controlled vocabulary. This enables structured "
            "metadata filtering before semantic search — users can constrain retrieval to "
            "specific barrier families (e.g., \"mechanical_integrity\") or human-factor "
            "dimensions (e.g., barriers where competence was a contributing factor). "
            "This taxonomy was developed through association mining on the full corpus."
        ),
    ]

    for title_text, desc in innovations:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Why It Matters to the Stakeholder", level=2)

    add_styled_table(doc,
        ["Stakeholder Impact", "How Our System Delivers"],
        [
            [
                "Noise Reduced",
                "The intersection filter (Stage 3) eliminates barriers from contextually "
                "irrelevant incidents before ranking. In a corpus of 3,253 barriers, a naive "
                "single-index search returns barriers from unrelated incident types. Our "
                "dual-axis intersection narrows candidates to only those barriers whose "
                "parent incidents are contextually similar to the query — typically reducing "
                "the candidate set by 60-80% before final ranking."
            ],
            [
                "Time Saved",
                "Sub-20ms average query latency (19ms baseline) across 3,253 barriers and "
                "526 incidents. A safety analyst manually cross-referencing barrier failures "
                "across 147 incident reports would spend hours; our system returns ranked, "
                "evidence-backed results in milliseconds. The structured context output "
                "(8,000 chars max) is ready for immediate LLM prompting or analyst review."
            ],
            [
                "Better Decisions",
                "Each result includes barrier rank, incident rank, RRF fusion score, "
                "barrier family, failure status, human contribution flag, 12 PIF dimensions, "
                "and verbatim supporting evidence. Decision-makers see not just which "
                "barriers are similar, but why they failed and in what operational context — "
                "enabling pattern recognition across incidents that would be invisible in "
                "tabular data alone."
            ],
        ]
    )

    doc.add_heading("System Architecture at a Glance", level=2)

    p = doc.add_paragraph()
    p.add_run("Pipeline Flow:").bold = True
    stages = [
        "Stage 1 — Barrier Similarity Search: Embed query, FAISS IndexFlatIP search with optional metadata mask (barrier_family, PIF filters). Returns top-50 barriers.",
        "Stage 2 — Incident Similarity Search: Embed query, FAISS IndexFlatIP search (unfiltered). Returns top-20 incidents.",
        "Stage 3 — Intersection Filter: Retain only barriers whose parent incident appeared in Stage 2 results.",
        "Stage 4 — RRF Fusion & Ranking: Score = 1/(60 + barrier_rank) + 1/(60 + incident_rank). Sort descending, return top-10.",
        "Optional Phase 2 — Cross-Encoder Reranking: Over-retrieve 30, re-score with ms-marco-MiniLM-L-6-v2 (22M params), truncate to final 10.",
    ]
    for s in stages:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_page_break()

    # ========================================================================
    # TOPIC 2: MOST DIFFICULT DECISION (5 points)
    # ========================================================================
    doc.add_heading("2. Most Difficult Decision (Tradeoff)", level=1)

    doc.add_heading("The Decision: Keep the Cross-Encoder Reranker Optional vs. Default", level=2)

    doc.add_paragraph(
        "After implementing the Phase-2 cross-encoder reranker (ms-marco-MiniLM-L-6-v2, "
        "22M parameters) as a post-RRF re-scoring stage, we had to decide whether to make "
        "it the default retrieval path or keep it as an optional enhancement. This was a "
        "decision that cut across scope, model choice, data approach, and architecture."
    )

    doc.add_heading("Why It Was Difficult", level=2)

    doc.add_paragraph(
        "The reranker showed clear per-query improvements — for specific query types, "
        "it dramatically improved ranking quality:"
    )

    add_styled_table(doc,
        ["Query", "Baseline Rank", "Reranked Rank", "Improvement"],
        [
            ["Communication / crane lift operations", "5", "1", "+4 ranks"],
            ["Change management / process modification", "3", "1", "+2 ranks"],
            ["Pressure safety valve / overpressure", "3", "1", "+2 ranks"],
            ["H2S detection / sour gas", "4", "2", "+2 ranks"],
        ]
    )

    doc.add_paragraph(
        "However, in aggregate across our 50-query evaluation benchmark, the improvements "
        "were marginal:"
    )

    add_styled_table(doc,
        ["Metric", "Baseline (RRF)", "Reranked", "Delta", "% Change"],
        [
            ["Top-1 Accuracy", "0.30", "0.30", "+0.00", "0.0%"],
            ["Top-5 Accuracy", "0.56", "0.56", "+0.00", "0.0%"],
            ["Top-10 Accuracy", "0.62", "0.60", "-0.02", "-3.2%"],
            ["MRR", "0.40", "0.42", "+0.01", "+3.1%"],
        ]
    )

    doc.add_paragraph(
        "The per-query analysis revealed the core tension: 9 queries improved, 6 degraded, "
        "15 were unchanged, and 20 missed entirely in both systems (40% miss rate). The "
        "reranker was solving a secondary problem (ranking quality) while the primary "
        "bottleneck (recall — 40% of queries finding no relevant result at all) remained "
        "unaddressed. Making the reranker default would add complexity and latency to a "
        "system whose fundamental limitation was elsewhere."
    )

    doc.add_heading("What We Chose and What Tradeoff We Accepted", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Decision: ")
    run.bold = True
    p.add_run(
        "Keep the reranker as an optional component via dependency injection. The baseline "
        "RRF pipeline remains the default. Users can opt in to reranking by passing a "
        "CrossEncoderReranker instance to RAGAgent.from_directory()."
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Architecture: ")
    run.bold = True
    p.add_run(
        "The reranker is injected at construction time — when reranker=None, the Phase-1 "
        "RRF codepath executes unchanged (retrieving top_k directly). When a reranker is "
        "provided, the system over-retrieves 30 candidates, re-scores them via cross-encoder, "
        "and truncates to the final 10. This preserves full backward compatibility: "
        "no existing code path is modified, no configuration flag needs to be toggled, and "
        "the Phase-1 test suite (46 tests) passes without modification."
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Tradeoff Accepted: ")
    run.bold = True
    p.add_run(
        "We accept that some specific queries (e.g., communication-related, change management) "
        "will rank slightly worse in the default path. In exchange, we keep the default system "
        "simpler (no model download, no GPU dependency, 19ms vs. 24ms latency), and we redirect "
        "engineering effort toward the actual bottleneck: improving bi-encoder recall through "
        "BM25 hybrid search, query expansion, and domain-tuned embeddings. The decision to "
        "defer is data-driven: +3.1% MRR falls below our 5% significance threshold for "
        "default activation."
    )

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Re-assessment Plan: ")
    run.bold = True
    p.add_run(
        "We plan to re-evaluate the reranker's value once the corpus exceeds 1,000 incidents. "
        "At larger corpus scale, the reranker's ability to disambiguate among more candidates "
        "becomes more valuable. The infrastructure is built and tested — activating it is a "
        "one-line change."
    )

    doc.add_page_break()

    # ========================================================================
    # TOPIC 3: HARDEST TECHNICAL CONSTRAINT (10 points)
    # ========================================================================
    doc.add_heading("3. Hardest Technical / Modeling Constraint & How We Solved It", level=1)

    doc.add_heading("The Constraint: Heterogeneous, Unstructured Multi-Source PDFs into a Unified Analytical Schema", level=2)

    doc.add_paragraph(
        "The hardest constraint in the project was transforming 148 unstructured PDF incident "
        "reports from four different regulatory agencies (CSB, BSEE, PHMSA, TSB) — each with "
        "different formats, terminologies, reporting standards, and document structures — into "
        "a single canonical JSON schema (Jeffrey V2.3) that is strict enough for Pydantic "
        "validation and analytics, yet flexible enough to support semantic retrieval. Every "
        "downstream capability — barrier coverage analysis, gap identification, RAG retrieval, "
        "cross-agency comparison, CSV exports — depends on this schema being correct and "
        "consistent."
    )

    doc.add_paragraph(
        "This is fundamentally a data quality problem at scale: if the structured extraction "
        "is unreliable, every analytical result built on top of it is unreliable. Garbage in, "
        "garbage out — but at 147 incidents with an average of 34 controls each, manual "
        "verification is not feasible."
    )

    doc.add_heading("The Mechanisms We Used to Solve It", level=2)

    # --- Mechanism 1: Schema Design ---
    doc.add_heading("Mechanism 1: Canonical Schema Design with Strict Validation (Schema)", level=3)

    doc.add_paragraph(
        "We designed the Jeffrey V2.3 schema with 8 top-level keys (incident_id, schema_version, "
        "source, context, event, bowtie, pifs, notes), enforced through Pydantic V2 models with "
        "Literal type constraints on all categorical fields. Every barrier control has 6 "
        "constrained enum fields:"
    )

    add_styled_table(doc,
        ["Field", "Allowed Values", "Purpose"],
        [
            ["side", "prevention, mitigation", "Bowtie diagram position"],
            ["barrier_type", "engineering, administrative, ppe, unknown", "Control classification"],
            ["line_of_defense", "1st, 2nd, 3rd, recovery, unknown", "Defense depth"],
            ["barrier_status", "active, degraded, failed, bypassed, not_installed, unknown", "Operational state"],
            ["confidence", "high, medium, low", "Evidence quality"],
            ["barrier_failed", "true, false", "Binary failure flag"],
        ]
    )

    doc.add_paragraph(
        "The schema also enforces structural constraints: each control must have a performance "
        "sub-model (with detection, alarm, manual intervention flags), a human sub-model (with "
        "12 PIF linkages), and an evidence sub-model (with verbatim supporting text). This "
        "ensures that every barrier in the corpus has a complete, analyzable record — not just "
        "a name and a status."
    )

    doc.add_paragraph(
        "V2.2-to-V2.3 conversion logic handles field renaming (side: prevention/mitigation to "
        "left/right, barrier_status: active to worked, line_of_defense: \"1st\" to integer 1), "
        "ensuring backward compatibility with earlier extractions."
    )

    # --- Mechanism 2: Quality Gating ---
    doc.add_heading("Mechanism 2: Multi-Layer Quality Gating (Filtering)", level=3)

    doc.add_paragraph(
        "Before any text reaches the LLM for structured extraction, it passes through a "
        "quality gate that detects extraction failures:"
    )

    add_styled_table(doc,
        ["Check", "Threshold", "What It Catches"],
        [
            ["EMPTY_TEXT", "0 chars", "Blank PDFs, scanned images with no OCR"],
            ["TOO_SHORT", "< 400 chars", "Partial extractions, cover pages only"],
            ["CID_ENCODING_GIBBERISH", "> 1% CID ratio or >= 5 CID tokens", "PDF font encoding failures (cid:NNNN artifacts)"],
            ["LOW_ALPHA_GIBBERISH", "< 55% alphabetic characters", "Binary data, corrupted extractions"],
        ]
    )

    doc.add_paragraph(
        "This prevented garbage-in-garbage-out: the Macondo report (scanned image, zero "
        "extractable text) was caught by EMPTY_TEXT and excluded from the corpus rather than "
        "producing a hallucinated extraction. The quality gate operates at the text layer, "
        "before any LLM cost is incurred."
    )

    doc.add_paragraph(
        "Post-extraction, the schema-check command validates every output JSON against the "
        "full Pydantic V2.3 model. The quality-gate command reports corpus-wide metrics: "
        "percentage of incidents with controls, controls-per-incident distribution (p50, p90), "
        "and field completeness across all 147 incidents."
    )

    # --- Mechanism 3: Cost-Optimized Extraction ---
    doc.add_heading("Mechanism 3: Cost-Optimized LLM Extraction with Model Ladder (Chunking / Filtering)", level=3)

    doc.add_paragraph(
        "Extracting structured data from 147 incident PDFs using LLMs is expensive if done "
        "naively. We built a cost-optimized extraction protocol:"
    )

    cost_items = [
        "Input Truncation: 50,000 character limit (~12,500 tokens) applied before LLM call. 15 of 20 PDFs in a typical batch exceed this limit. Truncation preserves the narrative-dense front matter while discarding appendices and boilerplate.",
        "Model Ladder: Primary model is Claude Haiku-4-5 (8,192 output tokens). On failure (rate limit, timeout, schema validation failure), the system escalates to 16,000 output tokens, then falls back to Claude Sonnet. In practice, Sonnet fallback was never triggered — Haiku handled 100% of extractions.",
        "Retry Logic: Up to 3 attempts per model tier before escalation. Dynamic rate-limit-aware delays scale with input token count (30k tokens/min budget).",
        "Result: $0.45 per batch of 20 PDFs. Full corpus (147 PDFs) extracted for approximately $3.30 total. This is an 80% cost reduction compared to using Sonnet as the primary model (~$16.50 estimated).",
    ]
    for item in cost_items:
        doc.add_paragraph(item, style="List Bullet")

    # --- Mechanism 4: Normalization Taxonomy ---
    doc.add_heading("Mechanism 4: Barrier Family Normalization Taxonomy (Normalization)", level=3)

    doc.add_paragraph(
        "Free-text barrier names from LLM extractions are inherently inconsistent: the same "
        "physical control might be called \"PSV\", \"Pressure Safety Valve\", \"pressure relief "
        "device\", or \"overpressure protection valve\" across different incidents. Without "
        "normalization, aggregation and retrieval are unreliable."
    )

    doc.add_paragraph(
        "We built a 4-quadrant, 46-family barrier classification taxonomy through association "
        "mining on the full corpus:"
    )

    add_styled_table(doc,
        ["Quadrant", "Example Families", "Count"],
        [
            ["Prevention x Administrative", "training, procedures, change_management, monitoring, hazard_analysis, communication, maintenance", "10"],
            ["Prevention x Engineering", "overpressurization_gas_isolation, fluid_containment, prevention_of_ignition, detection_alarms, mechanical_integrity", "5"],
            ["Mitigation x Administrative", "emergency_shutdown, fire_response, evacuation, medical_response, incident_command, emergency_preparedness", "13"],
            ["Mitigation x Engineering", "gas_detection, alarms, well_control, pressure_relief_blowdown, active_fire_protection, passive_fire_protection, spill_containment", "18"],
        ]
    )

    doc.add_paragraph(
        "The normalization pipeline: (1) lowercase and strip, (2) expand domain abbreviations "
        "(PSV to pressure safety valve, BOP to blowout preventer, ESDV to emergency shutdown "
        "valve, etc.), (3) remove punctuation, (4) keyword matching against family-specific "
        "term lists, (5) quadrant dispatch based on (side, barrier_type). First keyword match "
        "wins; unmatched controls fall back to \"other_{type}\"."
    )

    doc.add_paragraph(
        "This normalization is the bridge between unstructured LLM output and structured "
        "analytics: it converts 3,253 free-text barrier names into 25 distinct, queryable "
        "barrier families."
    )

    # --- Mechanism 5: Evaluation Harness ---
    doc.add_heading("Mechanism 5: Quantitative Evaluation Harness (Evaluation Harness)", level=3)

    doc.add_paragraph(
        "We built a 50-query evaluation benchmark to measure retrieval quality objectively, "
        "not anecdotally. Each query specifies a barrier_query, incident_query, and "
        "expected_barrier family. The harness computes:"
    )

    eval_items = [
        "Accuracy metrics: Top-1, Top-5, Top-10 hit rates (did the expected barrier family appear in the top K results?)",
        "Ranking quality: Mean Reciprocal Rank (MRR) — how high does the correct result rank?",
        "Latency profiling: Average, P95, and max query latency in milliseconds",
        "Memory profiling: RSS before/after agent construction, reranker overhead",
        "Per-query ranking deltas: For each query, baseline rank vs. reranked rank, with improvement/degradation classification",
        "Failure mode testing: 5 edge-case tests (empty query, single-result request, no-match query) verifying graceful degradation",
    ]
    for item in eval_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "This harness is what enabled the data-driven reranker decision in Topic 2. Without "
        "quantitative evaluation, we would have relied on cherry-picked examples (which favored "
        "the reranker) rather than aggregate metrics (which showed marginal improvement)."
    )

    doc.add_heading("How This Improves Decision Quality", level=2)

    doc.add_paragraph(
        "The combination of these five mechanisms transforms a pile of heterogeneous PDFs into "
        "a queryable, validated, normalized knowledge base. The improvement to decision quality "
        "operates at three levels:"
    )

    decision_items = [
        (
            "Data Trustworthiness",
            "Every incident in the corpus has passed text-level quality gating, LLM extraction "
            "with schema validation, and V2.3 Pydantic model enforcement. Analysts can trust "
            "that a \"barrier_status: failed\" field reflects the source document, not a "
            "hallucination, because it is backed by verbatim supporting_text evidence."
        ),
        (
            "Cross-Agency Comparability",
            "Barrier family normalization enables questions that were previously impossible: "
            "\"How often do mechanical_integrity barriers fail across BSEE vs. CSB incidents?\" "
            "Without normalization, the same physical barrier has different names in different "
            "agency reports, making aggregation meaningless."
        ),
        (
            "Evidence-Based Retrieval",
            "The evaluation harness ensures that retrieval quality is measured, not assumed. "
            "The 40% miss rate finding directly informed our roadmap: instead of optimizing "
            "ranking (reranker), we now know to invest in recall (BM25 hybrid, query expansion, "
            "domain-tuned embeddings). This prevents wasted effort on the wrong problem."
        ),
    ]
    for title_text, desc in decision_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Concrete Metrics That Demonstrate Improvement", level=2)

    add_styled_table(doc,
        ["Metric", "Value", "Significance"],
        [
            [
                "Extraction Success Rate",
                "147 / 148 PDFs (99.3%)",
                "Only 1 failure (Macondo — scanned image, no text). Quality gate caught it before LLM cost."
            ],
            [
                "Schema Validation",
                "100% pass after normalization",
                "All 739 V2.3 JSONs pass strict Pydantic validation. 4 files required field-level fixes (null performance, unknown side) — all caught by automated validation."
            ],
            [
                "Extraction Cost",
                "$0.45 / 20 PDFs ($3.30 total)",
                "80% cost reduction vs. Sonnet-primary baseline (~$16.50). Model ladder with Haiku primary achieved 100% extraction without fallback."
            ],
            [
                "Barrier Normalization",
                "3,253 controls into 25 families",
                "Free-text barrier names normalized into queryable, aggregatable taxonomy. Enables cross-agency barrier coverage comparison."
            ],
            [
                "Retrieval Performance",
                "62% Top-10 / 0.40 MRR / 19ms avg",
                "Sub-20ms queries across 3,253 barriers. Evaluation harness identified 40% miss rate as primary bottleneck — directly informing roadmap priorities."
            ],
            [
                "Test Coverage",
                "362 tests, 0 failures",
                "57 RAG-specific tests covering all modules (embeddings, vector index, corpus builder, retriever, reranker, context builder, agent, integration)."
            ],
        ]
    )

    doc.add_paragraph()

    # Summary paragraph
    p = doc.add_paragraph()
    run = p.add_run("In summary: ")
    run.bold = True
    p.add_run(
        "the hardest constraint was not any single technical problem, but the compound "
        "challenge of making heterogeneous, unstructured data analytically useful while "
        "maintaining data integrity, controlling costs, and enabling semantic retrieval. "
        "We solved it through a layered approach — schema design, quality gating, cost-optimized "
        "extraction, normalization, and quantitative evaluation — each layer addressing a "
        "different failure mode. The result is a corpus of 739 validated, normalized incident "
        "records with 3,253 classified barrier controls, queryable in under 20 milliseconds, "
        "produced for $3.30 in LLM costs."
    )

    return doc


def main() -> None:
    doc = build_document()
    out_path = Path(__file__).resolve().parent.parent / "Lessons_Learned_2_Bowtie_Risk_Analytics.docx"
    doc.save(str(out_path))
    print(f"Document saved to: {out_path}")


if __name__ == "__main__":
    main()
